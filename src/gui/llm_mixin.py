"""LLM-обработка GUI, вынесенная из GigaTranscriberQtApp.

Mixin: методы работают со `self` главного окна (виджеты создаются в app_qt).
Поведение сохранено 1:1 — методы перенесены без изменений.
"""
from __future__ import annotations

import os
import shutil
import threading
from pathlib import Path

from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QFileDialog, QListWidgetItem, QMessageBox

from ..config import LLM_TEMPERATURE
from ..services import llm_service

SUMMARY_PROMPT = (
    "Ты аналитик встреч и голосовых сообщений. Сделай сильную, плотную и полезную выжимку транскрипта на русском языке. "
    "Убери повторы, слова-паразиты и шум распознавания. Сохрани только смысл. "
    "\n\nСтруктура ответа:"
    "\n1. Краткое резюме в 3-6 пунктах."
    "\n2. Ключевые договоренности и решения."
    "\n3. Важные факты, цифры, сроки, имена и роли — если они есть."
    "\n4. Риски, спорные места или открытые вопросы — если они есть."
    "\n\nПиши четко, по делу, без воды. Если часть информации в транскрипте неясна, пометь это явно и не выдумывай."
)

TASKS_PROMPT = (
    "Ты project manager assistant. Из транскрипта выдели только конкретные задачи и оформи их в максимально рабочем виде на русском языке. "
    "Игнорируй рассуждения, повторы и фоновые фразы. Не выдумывай задачи, которых нет в тексте. "
    "\n\nДля каждой задачи укажи:"
    "\n- Что нужно сделать"
    "\n- Кто ответственный / исполнитель, если это можно понять"
    "\n- Срок, дедлайн или ориентир по времени, если упомянут"
    "\n- Контекст или комментарий, если он важен"
    "\n- Приоритет, если он читается из разговора"
    "\n\nСначала дай список задач. Затем отдельным коротким блоком выведи: "
    "«Открытые вопросы / неясности». Если задач нет, напиши: «Явных задач не найдено»."
)


class LlmMixin:
    def _build_llm_prompt_text(self, transcript_text: str, prompt: str) -> str:
        return llm_service.build_prompt_text(transcript_text, prompt)

    def _run_llm_provider(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        raw = llm_settings.get("provider", "API")
        provider = "Other" if self._normalize_llm_provider(raw) == "Other" else raw
        try:
            return llm_service.run_provider(
                llm_settings, transcript_text, prompt,
                provider=provider, strict_empty_cli=True,
            )
        except llm_service.UnknownLLMProvider as exc:
            raise RuntimeError(self._t(
                f"Неизвестный LLM-провайдер: {exc.provider}",
                f"Unknown LLM provider: {exc.provider}",
            )) from exc

    def _run_llm_processing(self, llm_settings: dict, items: list, modes: list, export_formats: list):
        try:
            results = []
            total = len(items)
            last_name = "llm_result"
            provider = llm_settings.get("provider", "API")
            for item_index, item in enumerate(items, start=1):
                name = item["name"]
                last_name = name
                item_blocks = []
                for mode_suffix, mode_label, prompt in modes:
                    self.log(f"LLM: обработка {item_index}/{total} — {name} — {mode_label} — {provider}")
                    answer = self._run_llm_provider(llm_settings, item["text"], prompt)
                    saved_paths = self._save_llm_result(item, answer, mode_suffix, export_formats)
                    block = f"=== {name} / {mode_label} / {provider} ===\n{answer}"
                    if saved_paths:
                        block += "\n\nСохранено:\n" + "\n".join(saved_paths)
                    item_blocks.append(block)
                results.append("\n\n".join(item_blocks))
            mode_suffixes = "_".join(mode[0] for mode in modes)
            self.llm_last_result_name = f"{last_name}_llm_{mode_suffixes}"
            final_text = "\n\n".join(results)
            self.signals.llm_finished.emit(True, f"LLM-обработка завершена: {total} файл(ов)", final_text)
        except Exception as e:
            error_text = str(e).strip() or "Неизвестная ошибка"
            self.signals.llm_finished.emit(False, f"Ошибка LLM: {self._compact_llm_error(error_text)}", error_text)

    def _compact_llm_error(self, error_text: str, limit: int = 180) -> str:
        raw_text = (error_text or "").strip()
        lowered = raw_text.lower()

        friendly_rules = [
            (("refresh token was revoked", "please log out and sign in again"), "Codex: сессия истекла или токен отозван — нужно заново войти в Codex"),
            (("token_invalidated", "authentication token has been invalidated"), "Codex: токен недействителен — перелогиньтесь"),
            (("your session has ended", "refresh_token_invalidated"), "Codex: сессия завершилась — выполните codex logout и codex login"),
            (("connection refused", "127.0.0.1", "/v1/responses"), "Codex: локальный backend недоступен — проверьте, запущен ли нужный сервер/провайдер"),
            (("failed to refresh available models", "missing field `base_instructions`"), "Codex: сервер моделей отдает несовместимый формат ответа — провайдер/прокси не полностью совместим с Codex"),
            (("failed to decode models response",), "Codex: провайдер вернул неожиданный формат списка моделей"),
            (("401", "anthropic"), "Anthropic API: ошибка авторизации (401) — проверьте API key"),
            (("401", "openai"), "OpenAI-compatible API: ошибка авторизации (401) — проверьте API key"),
            (("401", "unauthorized"), "Ошибка авторизации (401) — проверьте ключ, токен или логин выбранного провайдера"),
            (("403", "forbidden"), "Доступ запрещен (403) — у аккаунта или ключа не хватает прав"),
            (("404",), "Endpoint не найден (404) — проверьте URL API и путь /v1/..."),
            (("429", "rate"), "Превышен лимит запросов (429) — попробуйте позже или смените тариф/провайдера"),
            (("insufficient_quota",), "Закончилась квота API — проверьте биллинг или лимиты"),
            (("model_not_found",), "Указанная модель не найдена — проверьте точное имя модели"),
            (("does not exist", "model"), "Указанная модель не существует у выбранного провайдера"),
            (("invalid x-api-key",), "Неверный Anthropic API key"),
            (("incorrect api key",), "Неверный API key"),
            (("could not resolve host",), "Не удалось найти хост — проверьте URL и интернет-соединение"),
            (("name or service not known",), "Не удалось найти сервер — проверьте адрес API"),
            (("max retries exceeded",), "Не удалось подключиться к API после нескольких попыток"),
            (("read timed out", "timeout"), "Сервер слишком долго отвечает — попробуйте позже или увеличьте timeout"),
            (("connection timed out",), "Таймаут соединения — сервер недоступен или отвечает слишком долго"),
            (("ssl", "certificate"), "Ошибка SSL-сертификата — проверьте HTTPS/сертификат сервера"),
            (("command not found",), "Не найдена команда CLI-провайдера — проверьте путь в настройках"),
            (("not found", "claude"), "Claude Code не найден — проверьте путь к команде claude"),
            (("not found", "codex"), "Codex не найден — проверьте путь к команде codex"),
            (("not found", "opencode"), "OpenCode не найден — проверьте путь к команде opencode"),
            (("not found", "pi"), "Pi не найден — проверьте путь к команде pi"),
        ]
        for needles, message in friendly_rules:
            if all(needle in lowered for needle in needles):
                return message

        text = " ".join(raw_text.split())
        if len(text) <= limit:
            return text
        return text[:limit - 1] + "…"

    def _write_llm_output_file(self, save_path: str, content: str, export_format: str):
        if export_format in ("txt", "md"):
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(content)
            return
        if export_format == "docx":
            from docx import Document
            doc = Document()
            for block in content.split("\n\n"):
                doc.add_paragraph(block)
            doc.save(save_path)

    def _save_llm_result(self, item: dict, answer: str, mode_suffix: str, export_formats: list) -> list[str]:
        source_path = item.get("source_path")
        if self.llm_output_dir:
            target_dir = self.llm_output_dir
        elif source_path:
            target_dir = os.path.dirname(source_path)
        else:
            target_dir = os.getcwd()
        os.makedirs(target_dir, exist_ok=True)
        saved_paths = []
        for export_format in export_formats:
            save_path = os.path.join(target_dir, f"{item['name']}_llm_{mode_suffix}.{export_format}")
            self._write_llm_output_file(save_path, answer, export_format)
            saved_paths.append(save_path)
        return saved_paths

    def _set_llm_buttons_enabled(self, enabled: bool):
        self.btn_llm_process.setEnabled(enabled)
        self.btn_llm_clear.setEnabled(enabled)

    def _on_llm_finished(self, success: bool, message: str, result_text: str):
        self.is_llm_processing = False
        self._set_llm_buttons_enabled(True)
        self.lbl_llm_status.setText(message)
        if result_text:
            self.llm_last_result_text = result_text
            self.txt_llm_result.setPlainText(result_text)
        elif not success:
            self.llm_last_result_text = ""
            self.txt_llm_result.setPlainText(message)
        if success:
            QMessageBox.information(self, self._t("Готово", "Done"), message)
        else:
            QMessageBox.warning(self, self._t("Ошибка", "Error"), message)

    def _export_llm_result(self, export_format: str):
        result_text = self.txt_llm_result.toPlainText().strip()
        if not result_text:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("Нет результата для экспорта", "There is no result to export."))
            return
        suffix = {"txt": "txt", "md": "md", "docx": "docx"}[export_format]
        initial_dir = self.llm_output_dir or self.output_dir or os.path.expanduser("~")
        default_name = f"{self.llm_last_result_name}.{suffix}"
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            f"Сохранить результат как {suffix.upper()}",
            os.path.join(initial_dir, default_name),
            f"{suffix.upper()} files (*.{suffix})"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(f".{suffix}"):
            save_path += f".{suffix}"
        try:
            if export_format in ("txt", "md"):
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(result_text)
            else:
                try:
                    from docx import Document
                except Exception:
                    QMessageBox.warning(self, self._t("Ошибка", "Error"), self._t("Для экспорта в DOCX установите пакет python-docx", "Install the python-docx package to export to DOCX."))
                    return
                doc = Document()
                for block in result_text.split("\n\n"):
                    doc.add_paragraph(block)
                doc.save(save_path)
            target_dir = os.path.dirname(save_path)
            if target_dir:
                self.llm_output_dir = target_dir
                self.user_settings.set_value("llm_output_dir", target_dir)
                self._update_llm_output_dir_label(target_dir)
            self.lbl_llm_status.setText(f"Результат экспортирован: {os.path.basename(save_path)}")
        except Exception as e:
            QMessageBox.warning(self, self._t("Ошибка", "Error"), self._t("Не удалось экспортировать результат: ", "Failed to export the result: ") + str(e))

    def _refresh_llm_files_list(self):
        if not hasattr(self, "llm_files_list"):
            return
        self.llm_files_list.clear()
        for path in self.transcript_files_for_llm:
            item = QListWidgetItem(os.path.basename(path))
            item.setToolTip(path)
            item.setData(Qt.ItemDataRole.UserRole, path)
            self.llm_files_list.addItem(item)
        has_files = bool(self.transcript_files_for_llm)
        self.llm_files_list.setVisible(has_files)
        self.llm_drop_hint.setVisible(not has_files)
        if has_files:
            self._size_list_to_contents(self.llm_files_list)
        c = self._colors()
        if has_files:
            text = self._t(f"Выбрано транскриптов: {len(self.transcript_files_for_llm)}", f"Selected transcripts: {len(self.transcript_files_for_llm)}")
            self.lbl_llm_files.setText(text)
            self.lbl_llm_files_count.setText(text)
            self.lbl_llm_files.setStyleSheet(self._transparent_label_style(c["text_sub"]))
            self.lbl_llm_files_count.setStyleSheet(self._transparent_label_style(c["text_sub"]))
        else:
            text = self._t("Файлы не выбраны", "No files selected")
            self.lbl_llm_files.setText(text)
            self.lbl_llm_files_count.setText(text)
            self.lbl_llm_files.setStyleSheet(self._transparent_label_style(c["text_mute"]))
            self.lbl_llm_files_count.setStyleSheet(self._transparent_label_style(c["text_mute"]))
        self._update_llm_files_controls()

    def _update_llm_files_controls(self):
        if not hasattr(self, "btn_clear_llm_files"):
            return
        has_files = bool(self.transcript_files_for_llm)
        self.btn_clear_llm_files.setEnabled(has_files and not self.is_llm_processing)
        self.btn_remove_llm_file.setEnabled(bool(self.llm_files_list.selectedItems()) and not self.is_llm_processing)

    def _remove_selected_llm_files(self):
        if self.is_llm_processing:
            return
        selected = {item.data(Qt.ItemDataRole.UserRole) for item in self.llm_files_list.selectedItems()}
        if not selected:
            return
        self.transcript_files_for_llm = [p for p in self.transcript_files_for_llm if p not in selected]
        self._refresh_llm_files_list()
        self.user_settings.set_value("last_selected_transcript_files", [p for p in self.transcript_files_for_llm if os.path.isfile(p)])

    def _clear_llm_files_list(self):
        if self.is_llm_processing or not self.transcript_files_for_llm:
            return
        self.transcript_files_for_llm = []
        self._refresh_llm_files_list()
        self.user_settings.set_value("last_selected_transcript_files", [])

    def _select_llm_transcript_files(self):
        initial_dir = self.user_settings.get_value("llm_transcript_dir", self.llm_transcript_dir)
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Выберите транскрипты",
            initial_dir,
            "Транскрипты (*.txt *.md *.srt *.vtt);;Текстовые файлы (*.txt *.md);;Все файлы (*.*)"
        )
        if files:
            self.transcript_files_for_llm = files
            folder = os.path.dirname(files[0])
            self.llm_transcript_dir = folder
            if not self.llm_output_dir:
                self.llm_output_dir = folder
                self._update_llm_output_dir_label(folder)
            self.user_settings.set_value("llm_transcript_dir", folder)
            self.user_settings.set_value("last_selected_transcript_files", files)
            self._refresh_llm_files_list()
            self.lbl_llm_status.setText(self._t("Транскрипты готовы к LLM-обработке", "Transcripts are ready for LLM processing"))

    def _select_llm_output_folder(self):
        initial_dir = self.llm_output_dir or self.output_dir or os.path.expanduser("~")
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения LLM-результатов", initial_dir)
        if folder:
            self.llm_output_dir = folder
            self.user_settings.set_value("llm_output_dir", folder)
            self._update_llm_output_dir_label(folder)

    # ──────────────────────────────────────────────────────────────
    # Загрузка по ссылке
    # ──────────────────────────────────────────────────────────────

    def _clear_llm_result(self):
        if self.is_llm_processing:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("LLM-обработка уже выполняется", "LLM processing is already running."))
            return
        self.txt_llm_result.clear()
        self.llm_last_result_text = ""
        self.llm_last_result_name = "llm_result"
        self.lbl_llm_status.setText(self._t("Готово к LLM-обработке", "Ready for LLM processing"))

    def _clear_llm_all(self):
        if self.is_llm_processing:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("LLM-обработка уже выполняется", "LLM processing is already running."))
            return
        self.transcript_files_for_llm = []
        self.txt_llm_transcript.clear()
        self._refresh_llm_files_list()
        self._clear_llm_result()
        self.user_settings.set_value("last_selected_transcript_files", [])
        self.user_settings.set_value("llm_manual_transcript", "")

    def _selected_llm_modes(self):
        modes = []
        if self.llm_action_checkboxes["summary"].isChecked():
            prompt = self.txt_llm_summary_prompt.toPlainText().strip() or SUMMARY_PROMPT
            modes.append(("summary", "Выжимка", prompt))
        if self.llm_action_checkboxes["tasks"].isChecked():
            prompt = self.txt_llm_tasks_prompt.toPlainText().strip() or TASKS_PROMPT
            modes.append(("tasks", "Задачи", prompt))
        if self.llm_action_checkboxes["custom"].isChecked():
            custom_prompt = self.txt_llm_custom_prompt.toPlainText().strip()
            if not custom_prompt:
                raise ValueError("Для режима «Свой промпт» укажите пользовательский промпт в меню «Настройки → LLM API…»")
            modes.append(("custom", "Свой промпт", custom_prompt))
        if not modes:
            raise ValueError("Выберите хотя бы один чекбокс в блоке «Что делать»")
        return modes

    def _selected_llm_export_formats(self):
        formats = [key for key, cb in self.llm_export_checkboxes.items() if cb.isChecked()]
        if not formats:
            raise ValueError("Выберите хотя бы один формат сохранения результата")
        return formats

    def _start_llm_processing(self):
        if self.is_llm_processing:
            return
        try:
            llm_settings = self._collect_llm_settings()
            items = self._collect_llm_inputs()
            modes = self._selected_llm_modes()
            export_formats = self._selected_llm_export_formats()
        except ValueError as e:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), str(e))
            return

        self._save_ui_settings()
        self.is_llm_processing = True
        self._set_llm_buttons_enabled(False)
        self.lbl_llm_status.setText("Идет LLM-обработка...")
        self.txt_llm_result.clear()
        threading.Thread(
            target=self._run_llm_processing,
            args=(llm_settings, items, modes, export_formats),
            daemon=True,
        ).start()

    def _collect_llm_settings(self) -> dict:
        provider = self.combo_llm_provider.currentText().strip() or "API"
        api_url = self.entry_llm_api_url.text().strip()
        api_key = self.entry_llm_api_key.text().strip()
        model = self.entry_llm_model.text().strip()
        temperature_text = self.entry_llm_temperature.text().strip() or str(LLM_TEMPERATURE)
        try:
            temperature = float(temperature_text)
        except ValueError as exc:
            raise ValueError("Temperature должно быть числом") from exc
        if not 0 <= temperature <= 2:
            raise ValueError("Temperature должно быть в диапазоне 0..2")

        if provider == "API":
            if not api_url:
                raise ValueError("Укажите API URL")
            if not api_key:
                raise ValueError("Укажите API Key")
            if not model:
                raise ValueError("Укажите модель")
        elif provider == "Claude Code":
            claude_path = self.entry_llm_claude_path.text().strip() or "claude"
            if not (shutil.which(claude_path) or os.path.isfile(claude_path)):
                raise ValueError(f"Не найден Claude Code: {claude_path}")
        elif provider == "Codex":
            codex_path = self.entry_llm_codex_path.text().strip() or "codex"
            if not (shutil.which(codex_path) or os.path.isfile(codex_path)):
                raise ValueError(f"Не найден Codex: {codex_path}")
        elif provider == "OpenCode":
            opencode_path = self.entry_llm_opencode_path.text().strip() or "opencode"
            if not (shutil.which(opencode_path) or os.path.isfile(opencode_path)):
                raise ValueError(f"Не найден OpenCode: {opencode_path}")
        elif provider == "Pi":
            pi_path = self.entry_llm_pi_path.text().strip() or "pi"
            if not (shutil.which(pi_path) or os.path.isfile(pi_path)):
                raise ValueError(f"Не найден Pi: {pi_path}")
        elif self._normalize_llm_provider(provider) == "Other":
            other_path = self.entry_llm_other_path.text().strip()
            if not other_path:
                raise ValueError(self._t("Укажите команду для провайдера «Другое»", "Specify a command for the 'Other' provider"))
            if not (shutil.which(other_path) or os.path.isfile(other_path)):
                raise ValueError(f"Не найдена команда: {other_path}")
        return {
            "provider": provider,
            "api_url": api_url,
            "api_key": api_key,
            "model": model,
            "temperature": temperature,
            "claude_path": self.entry_llm_claude_path.text().strip() or "claude",
            "claude_args": self.entry_llm_claude_args.text().strip(),
            "codex_path": self.entry_llm_codex_path.text().strip() or "codex",
            "codex_args": self.entry_llm_codex_args.text().strip(),
            "opencode_path": self.entry_llm_opencode_path.text().strip() or "opencode",
            "opencode_args": self.entry_llm_opencode_args.text().strip(),
            "pi_path": self.entry_llm_pi_path.text().strip() or "pi",
            "pi_provider": self.entry_llm_pi_provider.text().strip(),
            "pi_args": self.entry_llm_pi_args.text().strip(),
            "other_path": self.entry_llm_other_path.text().strip(),
            "other_args": self.entry_llm_other_args.text().strip(),
        }

    def _collect_llm_inputs(self):
        manual_text = self.txt_llm_transcript.toPlainText().strip()
        items = []
        if manual_text:
            base_name = "manual_transcript"
            if self.transcript_files_for_llm:
                base_name = Path(self.transcript_files_for_llm[0]).stem
            items.append({"name": base_name, "text": manual_text, "source_path": self.transcript_files_for_llm[0] if self.transcript_files_for_llm else None})
        for path in self.transcript_files_for_llm:
            try:
                with open(path, encoding="utf-8") as f:
                    text = f.read().strip()
            except OSError:
                continue
            if text:
                items.append({"name": Path(path).stem, "text": text, "source_path": path})
        if not items:
            raise ValueError("Выберите хотя бы один транскрипт или вставьте текст вручную")
        return items

    # ──────────────────────────────────────────────────────────────
    # Обработка файлов
    # ──────────────────────────────────────────────────────────────
