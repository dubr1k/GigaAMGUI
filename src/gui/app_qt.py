"""
Главное окно приложения GigaAM v3 Transcriber на PyQt6
Строгий профессиональный дизайн без ярких цветов
"""

import os
import json
import shlex
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from pathlib import Path

try:
    import fcntl
except ImportError:  # pragma: no cover - non-POSIX fallback
    fcntl = None

from PyQt6.QtCore import QByteArray, QEvent, QObject, Qt, QTimer, QUrl, pyqtSignal
from PyQt6.QtGui import (
    QAction,
    QColor,
    QDesktopServices,
    QDragEnterEvent,
    QDropEvent,
    QFont,
    QFontDatabase,
    QFontMetrics,
    QGuiApplication,
    QIcon,
    QKeySequence,
    QPalette,
)
from PyQt6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFrame,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QScrollArea,
    QSpinBox,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from ..config import (
    APP_TITLE, MEDIA_EXTENSIONS, OUTPUT_FORMATS, STATS_FILE,
    LLM_API_URL, LLM_API_KEY, LLM_MODEL, LLM_TEMPERATURE,
    save_env_value, user_config_dir,
)
from ..core import ModelLoader, TranscriptionProcessor
from ..utils import (
    AppLogger, AudioConverter, MediaDownloader, ProcessingStats,
    TimeFormatter, UserSettings, LLMClient, LLMSettings,
)

_BASE_FONT_PT = 12.0
_MIN_UI_SCALE = 0.85
_MAX_UI_SCALE = 1.75
_INSTANCE_LOCK_NAME = "instance.lock"
_OPEN_REQUESTS_NAME = "open_requests.jsonl"

SUMMARY_PROMPT = (
    "Ты аналитик встреч и голосовых сообщений. Сделай сильную, плотную и полезную выжимку транскрипта на русском языке. "
    "Убери повторы, слова-паразиты и шум распознавания. Сохрани только смысл. "
    "\n\nСтруктура ответа:" 
    "\n1. Краткое резюме в 3-6 пунктах." 
    "\n2. Ключевые договоренности и решения." 
    "\n3. Важные факты, цифры, сроки, имена и роли — если они есть." 
    "\n4. Риски, спорные места или открытые вопросы — если они есть." 
    "\n\nПиши четко, по делу, без воды. Если часть информации в транскрипте неясна, пометь это явно и не выдумывай."
)

TASKS_PROMPT = (
    "Ты project manager assistant. Из транскрипта выдели только конкретные задачи и оформи их в максимально рабочем виде на русском языке. "
    "Игнорируй рассуждения, повторы и фоновые фразы. Не выдумывай задачи, которых нет в тексте. "
    "\n\nДля каждой задачи укажи:" 
    "\n- Что нужно сделать" 
    "\n- Кто ответственный / исполнитель, если это можно понять" 
    "\n- Срок, дедлайн или ориентир по времени, если упомянут" 
    "\n- Контекст или комментарий, если он важен" 
    "\n- Приоритет, если он читается из разговора" 
    "\n\nСначала дай список задач. Затем отдельным коротким блоком выведи: " 
    "«Открытые вопросы / неясности». Если задач нет, напиши: «Явных задач не найдено»."
)


def _read_ui_scale() -> float:
    raw_value = os.getenv("GIGAAM_UI_SCALE", "1").strip().replace(",", ".")
    try:
        scale = float(raw_value)
    except ValueError:
        scale = 1.0
    return max(_MIN_UI_SCALE, min(_MAX_UI_SCALE, scale))


def _format_css_number(value: float) -> str:
    return f"{value:.2f}".rstrip("0").rstrip(".")


class WorkerSignals(QObject):
    """Сигналы для потока обработки"""
    log_message = pyqtSignal(str)
    progress_update = pyqtSignal(int)
    file_progress_update = pyqtSignal(int)
    current_file_info = pyqtSignal(str)
    processing_finished = pyqtSignal(bool, str)
    stage_update = pyqtSignal(str, float)
    download_progress = pyqtSignal(int)
    download_finished = pyqtSignal(list)
    download_failed = pyqtSignal(str)
    llm_finished = pyqtSignal(bool, str, str)


class GigaApplication(QApplication):
    """QApplication with macOS Finder/Dock open-file event support."""

    file_open_requested = pyqtSignal(list)

    def __init__(self, argv):
        super().__init__(argv)
        self._pending_open_paths = []

    def event(self, event):
        if event.type() == QEvent.Type.FileOpen:
            path = ""
            try:
                url = event.url()
                if url.isLocalFile():
                    path = url.toLocalFile()
            except Exception:
                path = ""
            if not path:
                try:
                    path = event.file()
                except Exception:
                    path = ""
            if path:
                self._pending_open_paths.append(path)
                self.file_open_requested.emit([path])
            return True
        return super().event(event)

    def take_pending_open_paths(self):
        paths = self._pending_open_paths[:]
        self._pending_open_paths.clear()
        return paths


class GigaTranscriberQtApp(QMainWindow):
    """Главное окно приложения для транскрибации на PyQt6"""

    def __init__(self):
        super().__init__()

        self.files_to_process = []
        self.output_dir = ""
        self.input_dir = ""
        self.is_processing = False
        self._last_generated_transcript_files = []
        self._cancel_requested = False
        self.start_time = None
        self.files_processed = 0
        self.total_files = 0
        self.file_estimates = {}
        self.total_estimated_time = 0
        self.time_spent = 0
        self.current_file_start_time = 0
        self.current_stage = None
        self.current_stage_progress = 0.0
        self._stage_start_time = 0.0
        self._current_filename = ""
        self.is_downloading = False
        self.start_processing_after_download = False
        self._last_result_dir = ""

        self.transcript_files_for_llm = []
        self.llm_output_dir = ""
        self.llm_transcript_dir = os.path.expanduser("~")
        self.is_llm_processing = False
        self.llm_last_result_text = ""
        self.llm_last_result_name = "llm_result"

        self.enable_diarization = False
        self.num_speakers = None
        self._diarization_prompt_open = False

        self.output_formats = {
            'txt': True,
            'txt_timecodes': True,
            'txt_diarize': False,
            'txt_diarize_timecodes': False,
            'md': False,
            'srt': False,
            'vtt': False,
        }

        self.app_logger = AppLogger()
        self.app_logger.log_session_start()
        self.model_loader = ModelLoader()
        self.stats = ProcessingStats(STATS_FILE)
        self.time_formatter = TimeFormatter()
        self.user_settings = UserSettings()
        self.media_downloader = MediaDownloader()

        self._theme = self.user_settings.settings.get("theme", "dark")
        self._lang = self.user_settings.settings.get("language", "ru")
        self._ui_scale = self._effective_ui_scale()

        self.signals = WorkerSignals()
        self.signals.log_message.connect(self._append_log)
        self.signals.progress_update.connect(self._update_total_progress)
        self.signals.file_progress_update.connect(self._update_file_progress)
        self.signals.current_file_info.connect(self._update_current_file_info)
        self.signals.processing_finished.connect(self._on_processing_finished)
        self.signals.stage_update.connect(self._on_stage_update)
        self.signals.download_progress.connect(self._update_download_progress)
        self.signals.download_finished.connect(self._on_download_finished)
        self.signals.download_failed.connect(self._on_download_failed)
        self.signals.llm_finished.connect(self._on_llm_finished)

        self.progress_timer = QTimer()
        self.progress_timer.timeout.connect(self._update_progress_display)

        saved_output_dir = self.user_settings.get_last_output_dir()
        saved_input_dir = self.user_settings.get_last_files_dir()
        saved_llm_output_dir = self.user_settings.get_value("llm_output_dir", "")
        saved_llm_transcript_dir = self.user_settings.get_value("llm_transcript_dir", "")

        if saved_output_dir:
            self.output_dir = saved_output_dir
        if saved_input_dir:
            self.input_dir = saved_input_dir
        if saved_llm_output_dir and os.path.isdir(saved_llm_output_dir):
            self.llm_output_dir = saved_llm_output_dir
        elif saved_output_dir:
            self.llm_output_dir = saved_output_dir
        if saved_llm_transcript_dir and os.path.isdir(saved_llm_transcript_dir):
            self.llm_transcript_dir = saved_llm_transcript_dir
        elif saved_input_dir:
            self.llm_transcript_dir = saved_input_dir

        self._init_ui()
        self._restore_ui_settings()
        self.setAcceptDrops(True)

        if saved_output_dir:
            self._update_output_dir_label(saved_output_dir)
        if saved_input_dir:
            self._update_input_dir_label(saved_input_dir)
        if self.llm_output_dir:
            self._update_llm_output_dir_label(self.llm_output_dir)

        self.app_logger.cleanup_old_logs()

    # ──────────────────────────────────────────────────────────────
    # Темы
    # ──────────────────────────────────────────────────────────────

    _LIGHT = {
        "bg":        "#f5f6f8",
        "bg_card":   "#ffffff",
        "border":    "#dde0e6",
        "text":      "#1a1a1a",
        "text_sub":  "#374151",
        "text_mute": "#9ca3af",
        "text_mute2":"#6b7280",
        "btn_bg":    "#ffffff",
        "btn_border":"#d1d5db",
        "btn_text":  "#374151",
        "btn_hover_bg":    "#f0f4ff",
        "btn_hover_border":"#3b82f6",
        "btn_hover_text":  "#1d4ed8",
        "accent":    "#3b82f6",
        "accent2":   "#2563eb",
        "accent3":   "#1d4ed8",
        "accent_dis":"#93c5fd",
        "clear_bg":  "#f3f4f6",
        "clear_text":"#6b7280",
        "clear_border":"#e5e7eb",
        "clear_hover_bg":"#fee2e2",
        "clear_hover_border":"#fca5a5",
        "clear_hover_text":"#dc2626",
        "input_bg":  "#ffffff",
        "input_sel": "#bfdbfe",
        "input_dis": "#f9fafb",
        "input_dis_text":"#9ca3af",
        "progress_bg":"#e5e7eb",
        "progress_chunk":"#3b82f6",
        "progress_chunk2":"#2563eb",
        "tab_bg":    "#e8eaee",
        "tab_text":  "#555",
        "tab_sel_bg":"#f5f6f8",
        "tab_sel_text":"#1a1a1a",
        "tab_accent":"#3b82f6",
        "tab_hover": "#dde1ea",
        "scroll_bg": "#f1f2f4",
        "scroll_handle":"#cbd5e1",
        "scroll_handle_hover":"#94a3b8",
        "status_bg": "#f3f4f6",
        "status_text":"#374151",
        "theme_btn": "🌙",
    }

    _DARK = {
        "bg":        "#1e1e21",
        "bg_card":   "#2d2d30",
        "border":    "#3e3e42",
        "text":      "#e8e8e8",
        "text_sub":  "#c8c8c8",
        "text_mute": "#6b6b6b",
        "text_mute2":"#888888",
        "btn_bg":    "#3a3a3d",
        "btn_border":"#4a4a4e",
        "btn_text":  "#d0d0d0",
        "btn_hover_bg":    "#45455a",
        "btn_hover_border":"#5b7ee5",
        "btn_hover_text":  "#a8c4ff",
        "accent":    "#4f7de8",
        "accent2":   "#3a6ad4",
        "accent3":   "#2c57be",
        "accent_dis":"#2c3f6b",
        "clear_bg":  "#35353a",
        "clear_text":"#888888",
        "clear_border":"#3e3e44",
        "clear_hover_bg":"#4a2020",
        "clear_hover_border":"#8b3a3a",
        "clear_hover_text":"#e05050",
        "input_bg":  "#252528",
        "input_sel": "#1a3a6b",
        "input_dis": "#222225",
        "input_dis_text":"#555558",
        "progress_bg":"#303035",
        "progress_chunk":"#4f7de8",
        "progress_chunk2":"#3a6ad4",
        "tab_bg":    "#2a2a2d",
        "tab_text":  "#909090",
        "tab_sel_bg":"#1e1e21",
        "tab_sel_text":"#e8e8e8",
        "tab_accent":"#4f7de8",
        "tab_hover": "#333338",
        "scroll_bg": "#252528",
        "scroll_handle":"#4a4a50",
        "scroll_handle_hover":"#6a6a72",
        "status_bg": "#252528",
        "status_text":"#c0c0c0",
        "theme_btn": "☀️",
    }

    def _colors(self):
        return self._DARK if self._theme == "dark" else self._LIGHT

    def _effective_ui_scale(self) -> float:
        app = QApplication.instance()
        font_pt = app.font().pointSizeF() if app else _BASE_FONT_PT
        if font_pt <= 0:
            font_pt = _BASE_FONT_PT
        font_scale = max(_MIN_UI_SCALE, min(_MAX_UI_SCALE, font_pt / _BASE_FONT_PT))
        ui_scale = font_scale * _read_ui_scale()
        return round(max(_MIN_UI_SCALE, min(_MAX_UI_SCALE, ui_scale)), 4)

    def _px(self, value: int | float) -> int:
        return max(1, int(round(value * self._ui_scale)))

    def _pt(self, value: int | float) -> float:
        return round(value * self._ui_scale, 2)

    def _pt_css(self, value: int | float) -> str:
        return _format_css_number(self._pt(value))

    def _transparent_label_style(
        self,
        color: str,
        font_pt: int | float | None = None,
        font_weight: str | None = None,
    ) -> str:
        parts = ["background: transparent", f"color: {color}"]
        if font_pt is not None:
            parts.append(f"font-size: {self._pt_css(font_pt)}pt")
        if font_weight:
            parts.append(f"font-weight: {font_weight}")
        return "; ".join(parts) + ";"

    def _font(self, point_size: int | float, weight: QFont.Weight = QFont.Weight.Normal, fixed: bool = False) -> QFont:
        font_kind = QFontDatabase.SystemFont.FixedFont if fixed else QFontDatabase.SystemFont.GeneralFont
        font = QFontDatabase.systemFont(font_kind)
        font.setPointSizeF(self._pt(point_size))
        font.setWeight(weight)
        return font

    def _tab_min_width(self, labels: tuple[str, ...]) -> int:
        metrics = QFontMetrics(self._font(11, QFont.Weight.Bold))
        text_width = max(metrics.horizontalAdvance(label) for label in labels)
        return text_width + self._px(36)

    def _label_column_width(self, labels: tuple[str, ...], extra: int = 6) -> int:
        """Ширина колонки, достаточная для самого длинного лейбла из набора.

        Используется, чтобы поля формы (значения справа от лейблов) всегда
        начинались с одной и той же координаты X, даже если тексты лейблов
        разной длины ("Провайдер:", "Модель:", "API URL:" и т.д.).
        """
        metrics = QFontMetrics(self._font(10))
        text_width = max(metrics.horizontalAdvance(label) for label in labels)
        return text_width + self._px(extra)

    def _form_label(self, text: str, column_width: int) -> QLabel:
        lbl = QLabel(text)
        lbl.setFixedWidth(column_width)
        lbl.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        return lbl

    def _toggle_theme(self):
        self._theme = "dark" if self._theme == "light" else "light"
        self.user_settings.settings["theme"] = self._theme
        self.user_settings._save_settings()
        self._apply_theme()
        self._btn_theme.setText(self._colors()["theme_btn"])

    def _toggle_language(self):
        self._lang = "en" if self._lang == "ru" else "ru"
        self.user_settings.settings["language"] = self._lang
        self.user_settings._save_settings()
        self._apply_language()

    def _t(self, ru: str, en: str) -> str:
        return ru if self._lang == "ru" else en

    def _normalize_llm_provider(self, provider: str) -> str:
        return "Other" if provider in {"Другое", "Other"} else provider

    def _apply_language(self):
        is_ru = self._lang == "ru"
        self._btn_lang.setText("EN" if is_ru else "RU")
        self.setWindowTitle(APP_TITLE if is_ru else "GigaAM v3 Transcriber")
        if hasattr(self, "_title_label"):
            self._title_label.setText("GigaAM v3: Транскрибация" if is_ru else "GigaAM v3: Transcription")
        if hasattr(self, "tabs"):
            self.tabs.setTabText(0, "Обработка" if is_ru else "Process")
            self.tabs.setTabText(1, "LLM")
            self.tabs.setTabText(2, "Журнал обработки" if is_ru else "Processing log")
        if hasattr(self, "btn_start"):
            self.btn_start.setText("ЗАПУСТИТЬ ОБРАБОТКУ" if is_ru else "START PROCESSING")
        if hasattr(self, "btn_clear"):
            self.btn_clear.setText("ОЧИСТИТЬ ВСЕ" if is_ru else "CLEAR ALL")
        if hasattr(self, "btn_llm_process"):
            self.btn_llm_process.setText("ОБРАБОТАТЬ" if is_ru else "PROCESS")
        if hasattr(self, "btn_llm_clear"):
            self.btn_llm_clear.setText("ОЧИСТИТЬ ВСЕ" if is_ru else "CLEAR ALL")
        if hasattr(self, "status_bar"):
            self.status_bar.showMessage("Готов к работе" if is_ru else "Ready to work")
        if hasattr(self, "grp_files"):
            self.grp_files.setTitle("1. Выбор файлов" if is_ru else "1. File selection")
            self.grp_output.setTitle("2. Папка сохранения результатов" if is_ru else "2. Output folder")
            self.grp_diarization.setTitle("3. Диаризация спикеров" if is_ru else "3. Speaker diarization")
            self.grp_formats.setTitle("4. Форматы вывода" if is_ru else "4. Output formats")
            self.lbl_overall.setText("Общий прогресс" if is_ru else "Overall progress")
            self.btn_select_files.setText("Выбрать файлы" if is_ru else "Choose files")
            self.btn_select_files.setToolTip("Выбрать аудио/видео файлы для обработки  (Ctrl+O)" if is_ru else "Choose audio/video files for processing  (Ctrl+O)")
            self.btn_select_folder.setText("Выбрать папку" if is_ru else "Choose folder")
            self.btn_select_folder.setToolTip("Добавить все медиафайлы из папки и подпапок" if is_ru else "Add all media files from the folder and subfolders")
            self.btn_upload.setText("Загрузить" if is_ru else "Download")
            self.btn_upload.setToolTip("Скачать медиа по ссылке и добавить в очередь" if is_ru else "Download media by URL and add it to the queue")
            self.btn_output_select.setText("Выбрать папку" if is_ru else "Choose folder")
            self.btn_open_result.setText("Открыть папку с результатами" if is_ru else "Open results folder")
            if hasattr(self, "lbl_output_folder") and (self.lbl_output_folder.text().startswith("Папка не выбрана") or self.lbl_output_folder.text().startswith("Folder not selected")):
                self.lbl_output_folder.setText("Папка не выбрана (по умолчанию - рядом с файлом)" if is_ru else "Folder not selected (default: next to the file)")
            self.btn_cancel.setText("Отменить" if is_ru else "Cancel")
            self.cb_diarization.setText("Включить диаризацию спикеров" if is_ru else "Enable speaker diarization")
            self.cb_diarization.setToolTip("Определять, кто из спикеров говорит (нужен HF_TOKEN)" if is_ru else "Detect which speaker is talking (HF_TOKEN required)")
            self.lbl_num_speakers.setText("Кол-во спикеров:" if is_ru else "Speakers count:")
            self.lbl_diarization_info.setText("Автоматическое определение спикеров (требуется HF_TOKEN)" if is_ru else "Automatic speaker detection (HF_TOKEN required)")
            self.entry_num_speakers.setSpecialValueText("Авто" if is_ru else "Auto")
            self.entry_num_speakers.setToolTip("0 = автоопределение количества спикеров" if is_ru else "0 = auto-detect speaker count")
            self.input_path.setPlaceholderText("Ссылка на медиа (YouTube и др.)" if is_ru else "Media URL (YouTube, etc.)")
            self.input_path.setToolTip("Вставьте ссылку и нажмите «Загрузить»" if is_ru else "Paste a link and press 'Download'")
            if not self.files_to_process:
                self.lbl_files_count.setText("Файлы не выбраны" if is_ru else "No files selected")
            self.btn_remove_file.setText("Убрать выбранное" if is_ru else "Remove selected")
            self.btn_remove_file.setToolTip("Убрать выделенные файлы из очереди  (Delete)" if is_ru else "Remove selected files from the queue  (Delete)")
            self.btn_clear_files.setText("Очистить список" if is_ru else "Clear list")
            self.btn_clear_files.setToolTip("Убрать все файлы из очереди (настройки сохранятся)" if is_ru else "Remove all files from the queue (settings will be kept)")
            self.files_list.setToolTip("Очередь файлов. Выделите и нажмите Delete, чтобы убрать." if is_ru else "File queue. Select items and press Delete to remove them.")
            if self.lbl_input_folder.text().startswith("Папка не выбрана") or self.lbl_input_folder.text().startswith("Folder not selected"):
                self.lbl_input_folder.setText("Папка не выбрана" if is_ru else "Folder not selected")
            self.drop_hint.setText("Перетащите сюда файлы или папки  ·  либо нажмите «Выбрать файлы»" if is_ru else "Drop files or folders here  ·  or click 'Choose files'")
            format_labels = {
                "txt": ("Текст (.txt)", "Text (.txt)"),
                "txt_timecodes": ("Таймкоды (_timecodes.txt)", "Timecodes (_timecodes.txt)"),
                "txt_diarize": ("Диаризация (_diarize.txt)", "Diarization (_diarize.txt)"),
                "txt_diarize_timecodes": ("Диар.+тайм. (_diarize_timecodes.txt)", "Diarization+timecodes (_diarize_timecodes.txt)"),
                "md": ("Markdown (.md)", "Markdown (.md)"),
                "srt": ("SRT (.srt)", "SRT (.srt)"),
                "vtt": ("VTT (.vtt)", "VTT (.vtt)"),
            }
            for fmt, cb in self.format_checkboxes.items():
                ru_label, en_label = format_labels.get(fmt, (cb.text(), cb.text()))
                cb.setText(ru_label if is_ru else en_label)
        if hasattr(self, "grp_llm_source"):
            self.grp_llm_source.setTitle("1. Источник транскрипта" if is_ru else "1. Transcript source")
            self.grp_llm_output.setTitle("2. Куда сохранить" if is_ru else "2. Save location")
            self.grp_llm_actions.setTitle("3. Что сделать" if is_ru else "3. What to do")
            self.grp_llm_save.setTitle("4. Форматы вывода" if is_ru else "4. Output formats")
            self.grp_llm_result.setTitle("5. Результат LLM" if is_ru else "5. LLM result")
            self.btn_select_transcripts.setText("Выбрать транскрипты" if is_ru else "Choose transcripts")
            self.btn_llm_output.setText("Выбрать папку" if is_ru else "Choose folder")
            self.btn_llm_process.setToolTip("Запустить LLM-обработку выбранных транскриптов" if is_ru else "Run LLM processing for selected transcripts")
            self.btn_llm_clear.setToolTip("Сбросить выбранные транскрипты, ручной текст и результат LLM" if is_ru else "Reset selected transcripts, manual text and LLM result")
            if hasattr(self, "lbl_llm_summary_prompt"):
                self.lbl_llm_summary_prompt.setText("Промпт для выжимки:" if is_ru else "Prompt for summary:")
            if hasattr(self, "lbl_llm_tasks_prompt"):
                self.lbl_llm_tasks_prompt.setText("Промпт для задач:" if is_ru else "Prompt for tasks:")
            self.lbl_llm_supported.setText("Поддерживаемые файлы: .txt, .md, .srt, .vtt — либо вставьте транскрипт вручную ниже" if is_ru else "Supported files: .txt, .md, .srt, .vtt — or paste the transcript manually below")
            self.lbl_llm_status.setText("Готово к LLM-обработке" if is_ru else "Ready for LLM processing")
            if hasattr(self, "llm_drop_hint"):
                self.llm_drop_hint.setText("Перетащите сюда транскрипты  ·  либо нажмите «Выбрать транскрипты»" if is_ru else "Drop transcripts here  ·  or click 'Choose transcripts'")
            if hasattr(self, "btn_remove_llm_file"):
                self.btn_remove_llm_file.setText("Убрать выбранное" if is_ru else "Remove selected")
            if hasattr(self, "btn_clear_llm_files"):
                self.btn_clear_llm_files.setText("Очистить список" if is_ru else "Clear list")
            if hasattr(self, "llm_files_list"):
                self.llm_files_list.setToolTip("Список транскриптов. Выделите и нажмите Delete, чтобы убрать." if is_ru else "Transcript list. Select items and press Delete to remove them.")
            self.lbl_llm_files.setText("Файлы не выбраны" if is_ru and not self.transcript_files_for_llm else ("No files selected" if not is_ru and not self.transcript_files_for_llm else self.lbl_llm_files.text()))
            if hasattr(self, "lbl_llm_files_count") and not self.transcript_files_for_llm:
                self.lbl_llm_files_count.setText("Файлы не выбраны" if is_ru else "No files selected")
            self.txt_llm_transcript.setPlaceholderText(
                "Вставьте сюда транскрипт, если не хотите выбирать файлы"
                if is_ru else
                "Paste transcript here if you do not want to choose files"
            )
            if hasattr(self, "llm_action_checkboxes"):
                self.llm_action_checkboxes["summary"].setText("Выжимка" if is_ru else "Summary")
                self.llm_action_checkboxes["tasks"].setText("Задачи" if is_ru else "Tasks")
                self.llm_action_checkboxes["custom"].setText("Свой промпт" if is_ru else "Custom prompt")
            if hasattr(self, "lbl_llm_actions_note"):
                self.lbl_llm_actions_note.setText("Отметьте один или несколько режимов обработки. Для «Свой промпт» текст задается в меню «Настройки → LLM API…»." if is_ru else "Select one or more processing modes. For 'Custom prompt', set the text in Settings → LLM API…")
            if hasattr(self, "lbl_llm_output") and (self.lbl_llm_output.text().startswith("Папка не выбрана") or self.lbl_llm_output.text().startswith("Folder not selected")):
                self.lbl_llm_output.setText("Папка не выбрана (по умолчанию - рядом с транскриптом)" if is_ru else "Folder not selected (default: next to the transcript)")
            if hasattr(self, "lbl_llm_output_note"):
                self.lbl_llm_output_note.setText("Если папка не выбрана, результат будет сохранен рядом с исходным транскриптом." if is_ru else "If no folder is selected, the result will be saved next to the source transcript.")
            if hasattr(self, "llm_export_checkboxes"):
                self.llm_export_checkboxes["txt"].setText("TXT (.txt)")
                self.llm_export_checkboxes["md"].setText("Markdown (.md)")
                self.llm_export_checkboxes["docx"].setText("DOCX (.docx)")
            if hasattr(self, "btn_log_copy"):
                self.btn_log_copy.setText("Копировать" if is_ru else "Copy")
                self.btn_log_copy.setToolTip("Скопировать весь журнал в буфер обмена" if is_ru else "Copy the entire log to the clipboard")
                self.btn_log_save.setText("Сохранить…" if is_ru else "Save…")
                self.btn_log_save.setToolTip("Сохранить журнал в текстовый файл" if is_ru else "Save the log to a text file")
                self.btn_log_clear.setText("Очистить журнал" if is_ru else "Clear log")
                self.btn_log_clear.setToolTip("Очистить только журнал, не сбрасывая настройки" if is_ru else "Clear only the log without resetting settings")
            if hasattr(self, "_llm_settings_dialog"):
                self._llm_settings_dialog.setWindowTitle("Настройки LLM" if is_ru else "LLM settings")
                self.grp_llm_api_settings.setTitle("LLM API")
                self.llm_provider_labels["provider"].setText("Провайдер:" if is_ru else "Provider:")
                self.llm_provider_labels["model"].setText("Модель:" if is_ru else "Model:")
                self.llm_provider_items[5] = "Другое" if is_ru else "Other"
                current_provider = self._normalize_llm_provider(self.combo_llm_provider.currentText())
                self.combo_llm_provider.blockSignals(True)
                self.combo_llm_provider.setItemText(5, self.llm_provider_items[5])
                self.combo_llm_provider.setCurrentText(self.llm_provider_items[5] if current_provider == "Other" else current_provider)
                self.combo_llm_provider.blockSignals(False)
                self.llm_provider_labels["claude_path"].setText("Claude Code путь:" if is_ru else "Claude Code path:")
                self.llm_provider_labels["claude_args"].setText("Claude доп. аргументы:" if is_ru else "Claude extra args:")
                self.entry_llm_claude_args.setPlaceholderText("например: --permission-mode bypassPermissions" if is_ru else "example: --permission-mode bypassPermissions")
                self.llm_provider_labels["codex_path"].setText("Codex путь:" if is_ru else "Codex path:")
                self.llm_provider_labels["codex_args"].setText("Codex доп. аргументы:" if is_ru else "Codex extra args:")
                self.entry_llm_codex_args.setPlaceholderText("например: --dangerously-bypass-approvals-and-sandbox" if is_ru else "example: --dangerously-bypass-approvals-and-sandbox")
                self.llm_provider_labels["opencode_path"].setText("OpenCode путь:" if is_ru else "OpenCode path:")
                self.llm_provider_labels["opencode_args"].setText("OpenCode доп. аргументы:" if is_ru else "OpenCode extra args:")
                self.entry_llm_opencode_args.setPlaceholderText("например: --print" if is_ru else "example: --print")
                self.llm_provider_labels["pi_path"].setText("Pi путь:" if is_ru else "Pi path:")
                self.llm_provider_labels["pi_provider"].setText("Pi provider:" if is_ru else "Pi provider:")
                self.llm_provider_labels["pi_args"].setText("Pi доп. аргументы:" if is_ru else "Pi extra args:")
                self.entry_llm_pi_args.setPlaceholderText("например: --no-tools --thinking low" if is_ru else "example: --no-tools --thinking low")
                self.llm_provider_labels["other_path"].setText("Команда:" if is_ru else "Command:")
                self.llm_provider_labels["other_args"].setText("Аргументы:" if is_ru else "Arguments:")
                self.entry_llm_other_path.setPlaceholderText("путь к CLI, например my-llm" if is_ru else "CLI path, for example my-llm")
                self.entry_llm_other_args.setPlaceholderText("аргументы; промпт будет добавлен в конец как последний параметр" if is_ru else "arguments; the prompt will be appended as the last argument")
                self.prompts_group.setTitle("Готовые промпты" if is_ru else "Ready prompts")
                self.lbl_llm_summary_prompt.setText("Промпт для выжимки:" if is_ru else "Prompt for summary:")
                self.lbl_llm_tasks_prompt.setText("Промпт для задач:" if is_ru else "Prompt for tasks:")
                self.lbl_llm_settings_note.setText("Можно использовать OpenAI-compatible API, Anthropic Messages API, а также локальные Claude Code / Codex / OpenCode / Pi. Для API режим сам определяет тип API по URL или endpoint. Выбранный провайдер, модель, temperature, чекбоксы, prompt и файлы сохраняются между запусками. API Key лучше хранить в .env." if is_ru else "You can use an OpenAI-compatible API, Anthropic Messages API, or local Claude Code / Codex / OpenCode / Pi. In API mode, the app auto-detects the API type from the URL or endpoint. The selected provider, model, temperature, checkboxes, prompts, and files are saved between launches. It is best to store the API key in .env.")
                self._llm_settings_buttons.button(QDialogButtonBox.StandardButton.Save).setText("Сохранить" if is_ru else "Save")
                self._llm_settings_buttons.button(QDialogButtonBox.StandardButton.Close).setText("Закрыть" if is_ru else "Close")
                self._update_llm_provider_fields(self.combo_llm_provider.currentText())
        if hasattr(self, "_menu_file"):
            self._menu_file.setTitle("Файл" if is_ru else "File")
            self._menu_view.setTitle("Вид" if is_ru else "View")
            self._menu_settings.setTitle("Настройки" if is_ru else "Settings")
            self._menu_help.setTitle("Справка" if is_ru else "Help")
            self._act_files.setText("Выбрать файлы…" if is_ru else "Choose files…")
            self._act_folder.setText("Выбрать папку с файлами…" if is_ru else "Choose folder with files…")
            self._act_out.setText("Папка сохранения…" if is_ru else "Output folder…")
            self._act_open_res.setText("Открыть папку с результатами" if is_ru else "Open results folder")
            self._act_quit.setText("Выход" if is_ru else "Exit")
            self._act_theme.setText("Переключить тему" if is_ru else "Toggle theme")
            self._act_device.setText("Устройство (CPU / GPU)…" if is_ru else "Device (CPU / GPU)…")
            self._act_llm.setText("LLM API…")
            self._act_about.setText("О программе" if is_ru else "About")

    def _change_device(self):
        """Смена вычислительного устройства (CPU / GPU / GPU 50xx) из меню."""
        from .device_dialog import change_device_interactive
        from PyQt6.QtWidgets import QMessageBox

        if self.is_processing:
            QMessageBox.information(
                self,
                self._t("Устройство", "Device"),
                self._t("Дождитесь окончания обработки перед сменой устройства.", "Wait for processing to finish before changing the device."),
            )
            return

        changed = change_device_interactive(self)
        if changed:
            # Модель уже загружена под старую сборку torch — нужен перезапуск.
            self.model_loader.unload()

    def _apply_theme(self):
        c = self._colors()
        palette = QPalette()
        palette.setColor(QPalette.ColorRole.Window,          QColor(c["bg"]))
        palette.setColor(QPalette.ColorRole.WindowText,      QColor(c["text"]))
        palette.setColor(QPalette.ColorRole.Base,            QColor(c["input_bg"]))
        palette.setColor(QPalette.ColorRole.AlternateBase,   QColor(c["bg_card"]))
        palette.setColor(QPalette.ColorRole.Text,            QColor(c["text"]))
        palette.setColor(QPalette.ColorRole.Button,          QColor(c["btn_bg"]))
        palette.setColor(QPalette.ColorRole.ButtonText,      QColor(c["btn_text"]))
        palette.setColor(QPalette.ColorRole.Highlight,       QColor(c["accent"]))
        palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))
        self.setPalette(palette)

        r = c["progress_chunk"]
        r2 = c["progress_chunk2"]
        rad_f = self._px(11)
        rad_s = self._px(8)
        tab_min_width = self._tab_min_width(("Обработка", "Журнал обработки"))

        self.setStyleSheet(f"""
            QMainWindow, QWidget {{
                background-color: {c["bg"]};
                color: {c["text"]};
            }}
            QScrollArea {{
                background-color: {c["bg"]};
                border: none;
            }}
            QTabWidget::pane {{
                border: none;
                background-color: {c["bg"]};
            }}
            QTabBar::tab {{
                background-color: {c["tab_bg"]};
                color: {c["tab_text"]};
                border: 1px solid {c["border"]};
                border-bottom: none;
                border-radius: {self._px(5)}px {self._px(5)}px 0 0;
                padding: {self._px(6)}px {self._px(18)}px;
                font-size: {self._pt_css(11)}pt;
                margin-right: {self._px(2)}px;
                min-width: {tab_min_width}px;
            }}
            QTabBar::tab:selected {{
                background-color: {c["tab_sel_bg"]};
                color: {c["tab_sel_text"]};
                font-weight: bold;
                border-bottom: {self._px(2)}px solid {c["tab_accent"]};
            }}
            QTabBar::tab:hover:!selected {{
                background-color: {c["tab_hover"]};
            }}
            QGroupBox {{
                font-weight: bold;
                font-size: {self._pt_css(11)}pt;
                border: 1px solid {c["border"]};
                border-radius: {self._px(8)}px;
                margin-top: {self._px(18)}px;
                padding-top: {self._px(10)}px;
                background-color: {c["bg_card"]};
                color: {c["text"]};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top left;
                left: {self._px(14)}px;
                padding: 0 {self._px(6)}px 0 {self._px(6)}px;
                color: {c["text_sub"]};
                background: transparent;
            }}
            QPushButton {{
                background-color: {c["btn_bg"]};
                border: 1px solid {c["btn_border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(7)}px {self._px(17)}px {self._px(7)}px {self._px(17)}px;
                color: {c["btn_text"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QPushButton:hover {{
                background-color: {c["btn_hover_bg"]};
                border: 1px solid {c["btn_hover_border"]};
                color: {c["btn_hover_text"]};
            }}
            QPushButton:pressed {{
                background-color: {c["accent_dis"]};
                border: 1px solid {c["accent2"]};
            }}
            QPushButton:disabled {{
                background-color: {c["input_dis"]};
                color: {c["text_mute"]};
                border: 1px solid {c["border"]};
            }}
            QPushButton#start_button {{
                background-color: {c["accent"]};
                color: #ffffff;
                font-size: {self._pt_css(13)}pt;
                font-weight: bold;
                border: none;
                border-radius: {self._px(8)}px;
            }}
            QPushButton#start_button:hover {{
                background-color: {c["accent2"]};
            }}
            QPushButton#start_button:pressed {{
                background-color: {c["accent3"]};
            }}
            QPushButton#start_button:disabled {{
                background-color: {c["accent_dis"]};
                color: #ffffff;
            }}
            QPushButton#clear_button {{
                background-color: {c["clear_bg"]};
                color: {c["clear_text"]};
                font-size: {self._pt_css(10)}pt;
                font-weight: bold;
                border: 1px solid {c["clear_border"]};
                border-radius: {self._px(6)}px;
            }}
            QPushButton#clear_button:hover {{
                background-color: {c["clear_hover_bg"]};
                border: 1px solid {c["clear_hover_border"]};
                color: {c["clear_hover_text"]};
            }}
            QPushButton#theme_button {{
                background-color: transparent;
                border: 1px solid {c["border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(2)}px {self._px(8)}px;
                font-size: {self._pt_css(16)}pt;
                color: {c["text_sub"]};
            }}
            QPushButton#theme_button:hover {{
                background-color: {c["btn_hover_bg"]};
                border: 1px solid {c["btn_hover_border"]};
            }}
            QProgressBar {{
                border: none;
                border-radius: {rad_f}px;
                text-align: center;
                background-color: {c["progress_bg"]};
                color: {c["text"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QProgressBar::chunk {{
                background-color: qlineargradient(x1:0,y1:0,x2:1,y2:0,
                    stop:0 {r}, stop:1 {r2});
                border-radius: {rad_f}px;
            }}
            QLineEdit {{
                border: 1px solid {c["btn_border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(7)}px {self._px(11)}px {self._px(7)}px {self._px(11)}px;
                background-color: {c["input_bg"]};
                color: {c["text"]};
                selection-background-color: {c["input_sel"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QLineEdit:focus {{
                border: 1px solid {c["accent"]};
            }}
            QLineEdit:disabled {{
                background-color: {c["input_dis"]};
                color: {c["input_dis_text"]};
            }}
            QTextEdit {{
                border: 1px solid {c["border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(10)}px;
                background-color: {c["input_bg"]};
                color: {c["text"]};
                selection-background-color: {c["input_sel"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QCheckBox {{
                background: transparent;
                spacing: {self._px(8)}px;
                color: {c["text_sub"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QCheckBox::indicator {{
                width: {self._px(18)}px;
                height: {self._px(18)}px;
                border: 1.5px solid {c["btn_border"]};
                border-radius: {self._px(4)}px;
                background-color: {c["input_bg"]};
            }}
            QCheckBox::indicator:checked {{
                background-color: {c["accent"]};
                border: 1.5px solid {c["accent"]};
            }}
            QCheckBox::indicator:hover {{
                border: 1.5px solid {c["accent"]};
            }}
            QCheckBox:disabled {{
                color: {c["text_mute"]};
            }}
            QCheckBox::indicator:disabled {{
                background-color: {c["input_dis"]};
                border: 1.5px solid {c["border"]};
            }}
            QLabel {{
                background: transparent;
                color: {c["text_sub"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QScrollBar:vertical {{
                background: {c["scroll_bg"]};
                width: {self._px(8)}px;
                border-radius: {self._px(4)}px;
            }}
            QScrollBar::handle:vertical {{
                background: {c["scroll_handle"]};
                border-radius: {self._px(4)}px;
                min-height: {self._px(30)}px;
            }}
            QScrollBar::handle:vertical:hover {{
                background: {c["scroll_handle_hover"]};
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0px; }}
            QScrollBar:horizontal {{
                background: {c["scroll_bg"]};
                height: {self._px(8)}px;
                border-radius: {self._px(4)}px;
            }}
            QScrollBar::handle:horizontal {{
                background: {c["scroll_handle"]};
                border-radius: {self._px(4)}px;
                min-width: {self._px(30)}px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background: {c["scroll_handle_hover"]};
            }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{ width: 0px; }}
            #progress_card {{
                background-color: {c["bg_card"]};
                border: 1px solid {c["border"]};
                border-radius: {self._px(8)}px;
            }}
            #progress_card QLabel {{
                border: none;
                background: transparent;
            }}
            QListWidget#files_list, QListWidget#llm_files_list {{
                background-color: {c["input_bg"]};
                border: 1px solid {c["border"]};
                border-radius: {self._px(6)}px;
                color: {c["text"]};
                font-size: {self._pt_css(10)}pt;
                padding: {self._px(2)}px;
            }}
            QListWidget#files_list::item, QListWidget#llm_files_list::item {{
                padding: {self._px(3)}px {self._px(6)}px;
                border-radius: {self._px(4)}px;
            }}
            QListWidget#files_list::item:selected, QListWidget#llm_files_list::item:selected {{
                background-color: {c["accent"]};
                color: #ffffff;
            }}
            QListWidget#files_list::item:hover:!selected, QListWidget#llm_files_list::item:hover:!selected {{
                background-color: {c["btn_hover_bg"]};
            }}
            QSpinBox {{
                border: 1px solid {c["btn_border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(2)}px {self._px(8)}px;
                background-color: {c["input_bg"]};
                color: {c["text"]};
                selection-background-color: {c["input_sel"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QSpinBox:focus {{
                border: 1px solid {c["accent"]};
            }}
            QSpinBox:disabled {{
                background-color: {c["input_dis"]};
                color: {c["input_dis_text"]};
            }}
            QPushButton#cancel_button {{
                background-color: {c["clear_bg"]};
                color: {c["clear_text"]};
                font-size: {self._pt_css(9)}pt;
                font-weight: bold;
                border: 1px solid {c["clear_border"]};
                border-radius: {self._px(6)}px;
                padding: {self._px(2)}px {self._px(12)}px;
            }}
            QPushButton#cancel_button:hover {{
                background-color: {c["clear_hover_bg"]};
                border: 1px solid {c["clear_hover_border"]};
                color: {c["clear_hover_text"]};
            }}
            QPushButton#open_result_button {{
                background-color: transparent;
                color: {c["accent"]};
                font-size: {self._pt_css(10)}pt;
                font-weight: bold;
                border: 1px solid {c["accent"]};
                border-radius: {self._px(6)}px;
            }}
            QPushButton#open_result_button:hover {{
                background-color: {c["btn_hover_bg"]};
                border: 1px solid {c["btn_hover_border"]};
            }}
            QMenuBar {{
                background-color: {c["bg_card"]};
                color: {c["text_sub"]};
                border-bottom: 1px solid {c["border"]};
                font-size: {self._pt_css(10)}pt;
            }}
            QMenuBar::item {{
                background: transparent;
                padding: {self._px(4)}px {self._px(10)}px;
            }}
            QMenuBar::item:selected {{
                background-color: {c["btn_hover_bg"]};
                color: {c["btn_hover_text"]};
                border-radius: {self._px(4)}px;
            }}
            QMenu {{
                background-color: {c["bg_card"]};
                color: {c["text_sub"]};
                border: 1px solid {c["border"]};
                padding: {self._px(4)}px;
            }}
            QMenu::item {{
                padding: {self._px(5)}px {self._px(22)}px;
                border-radius: {self._px(4)}px;
            }}
            QMenu::item:selected {{
                background-color: {c["accent"]};
                color: #ffffff;
            }}
            QMenu::separator {{
                height: 1px;
                background: {c["border"]};
                margin: {self._px(4)}px {self._px(6)}px;
            }}
            QStatusBar {{
                background-color: {c["status_bg"]};
                color: {c["text_mute2"]};
                font-size: {self._pt_css(9)}pt;
            }}
            QToolTip {{
                background-color: {c["bg_card"]};
                color: {c["text"]};
                border: 1px solid {c["accent"]};
                padding: {self._px(4)}px;
            }}
        """)
        self._style_drop_hint()

        # Обновляем динамические стили прогресс-баров файла (тонкий)
        if hasattr(self, 'progress_bar_file'):
            self.progress_bar_file.setStyleSheet(
                f"QProgressBar {{ border: none; border-radius: {rad_s}px;"
                f"  background-color: {c['progress_bg']}; text-align: center;"
                f"  color: {c['text']}; font-size: {self._pt_css(8)}pt; font-weight: 600; }}"
                f"QProgressBar::chunk {{ border-radius: {rad_s}px;"
                f"  background-color: qlineargradient(x1:0,y1:0,x2:1,y2:0,"
                f"  stop:0 {r}, stop:1 {r2}); }}"
            )
        if hasattr(self, 'progress_bar_total'):
            self.progress_bar_total.setStyleSheet(
                f"QProgressBar {{ border: none; border-radius: {rad_f}px;"
                f"  background-color: {c['progress_bg']}; text-align: center;"
                f"  color: {c['text']}; font-size: {self._pt_css(10)}pt; font-weight: 600; }}"
                f"QProgressBar::chunk {{ border-radius: {rad_f}px;"
                f"  background-color: qlineargradient(x1:0,y1:0,x2:1,y2:0,"
                f"  stop:0 {r}, stop:1 {r2}); }}"
            )
        if hasattr(self, 'progress_upload'):
            self.progress_upload.setStyleSheet(
                f"QProgressBar {{ border: none; background-color: {c['progress_bg']};"
                f"  border-radius: {self._px(3)}px; }}"
                f"QProgressBar::chunk {{ background-color: {c['accent']}; border-radius: {self._px(3)}px; }}"
            )
        if hasattr(self, 'lbl_file_counter'):
            self.lbl_file_counter.setStyleSheet(
                f"color: {c['accent']}; font-size: {self._pt_css(11)}pt; font-weight: bold;"
            )
        if hasattr(self, 'lbl_status'):
            self.lbl_status.setStyleSheet(
                f"color: {c['status_text']}; font-size: {self._pt_css(10)}pt; font-weight: bold;"
                f"background-color: {c['status_bg']}; border-radius: {self._px(6)}px; padding: {self._px(2)}px;"
            )
        if hasattr(self, 'lbl_stage'):
            self.lbl_stage.setStyleSheet(
                self._transparent_label_style(c["text_sub"], font_pt=9, font_weight="600")
            )
        if hasattr(self, 'lbl_current_file'):
            self.lbl_current_file.setStyleSheet(self._transparent_label_style(c["text_mute2"], font_pt=9))

    # ──────────────────────────────────────────────────────────────
    # UI
    # ──────────────────────────────────────────────────────────────

    def _init_ui(self):
        self.setWindowTitle(APP_TITLE)
        self.setMinimumSize(self._px(940), self._px(680))
        self.resize(self._px(1040), self._px(1000))

        self._build_menu_bar()

        root = QWidget()
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(self._px(16), self._px(12), self._px(16), self._px(12))
        root_layout.setSpacing(self._px(8))
        self.setCentralWidget(root)

        # Заголовок + кнопка темы
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)

        self._header_left_spacer = QWidget()
        header_btn_width = self._px(48) + self._px(42) + header_row.spacing()
        self._header_left_spacer.setFixedWidth(header_btn_width)
        header_row.addWidget(self._header_left_spacer)

        self._title_label = QLabel("GigaAM v3: Транскрибация")
        self._title_label.setFont(self._font(18, QFont.Weight.Bold))
        self._title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._title_label.setFixedHeight(self._px(40))
        header_row.addWidget(self._title_label, 1)

        self._btn_lang = QPushButton("EN" if self._lang == "ru" else "RU")
        self._btn_lang.setObjectName("theme_button")
        self._btn_lang.setFixedSize(self._px(48), self._px(36))
        self._btn_lang.setToolTip("Switch language")
        self._btn_lang.clicked.connect(self._toggle_language)
        header_row.addWidget(self._btn_lang)

        self._btn_theme = QPushButton(self._colors()["theme_btn"])
        self._btn_theme.setObjectName("theme_button")
        self._btn_theme.setFixedSize(self._px(42), self._px(36))
        self._btn_theme.setToolTip("Переключить тему")
        self._btn_theme.clicked.connect(self._toggle_theme)
        header_row.addWidget(self._btn_theme)

        root_layout.addLayout(header_row)

        tabs = QTabWidget()
        tabs.tabBar().setElideMode(Qt.TextElideMode.ElideNone)
        root_layout.addWidget(tabs, 1)

        # ── Вкладка «Обработка» ──
        content_widget = QWidget()
        main_layout = QVBoxLayout(content_widget)
        main_layout.setContentsMargins(self._px(8), self._px(14), self._px(8), self._px(6))
        main_layout.setSpacing(self._px(4))

        main_layout.addWidget(self._create_files_group())
        main_layout.addWidget(self._create_output_group())
        main_layout.addWidget(self._create_diarization_group())
        main_layout.addWidget(self._create_formats_group())

        self.btn_start = QPushButton("ЗАПУСТИТЬ ОБРАБОТКУ")
        self.btn_start.setObjectName("start_button")
        self.btn_start.setFixedHeight(self._px(52))
        self.btn_start.setToolTip("Начать транскрибацию выбранных файлов  (Ctrl+Enter)")
        self.btn_start.setShortcut(QKeySequence("Ctrl+Return"))
        self.btn_start.clicked.connect(self._start_processing_thread)
        main_layout.addWidget(self.btn_start)

        self._create_progress_section(main_layout)

        self.btn_clear = QPushButton("ОЧИСТИТЬ ВСЕ")
        self.btn_clear.setObjectName("clear_button")
        self.btn_clear.setFixedHeight(self._px(40))
        self.btn_clear.setToolTip("Сбросить файлы, папки, журнал и прогресс")
        self.btn_clear.clicked.connect(self._clear_all)
        main_layout.addWidget(self.btn_clear)

        main_layout.addStretch()

        proc_scroll = QScrollArea()
        proc_scroll.setWidgetResizable(True)
        proc_scroll.setFrameShape(QFrame.Shape.NoFrame)
        proc_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        proc_scroll.setWidget(content_widget)
        tabs.addTab(proc_scroll, "Обработка")

        llm_scroll = QScrollArea()
        llm_scroll.setWidgetResizable(True)
        llm_scroll.setFrameShape(QFrame.Shape.NoFrame)
        llm_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        llm_scroll.setWidget(self._create_llm_tab())
        tabs.addTab(llm_scroll, "LLM")

        # ── Вкладка «Журнал» ──
        log_tab = QWidget()
        log_layout = QVBoxLayout(log_tab)
        log_layout.setContentsMargins(self._px(8), self._px(14), self._px(8), self._px(8))
        log_layout.setSpacing(self._px(6))

        log_toolbar = QHBoxLayout()
        log_toolbar.setSpacing(self._px(8))
        self.btn_log_copy = QPushButton("Копировать")
        self.btn_log_copy.setToolTip("Скопировать весь журнал в буфер обмена")
        self.btn_log_copy.setFixedHeight(self._px(32))
        self.btn_log_copy.clicked.connect(self._copy_log)
        log_toolbar.addWidget(self.btn_log_copy)
        self.btn_log_save = QPushButton("Сохранить…")
        self.btn_log_save.setToolTip("Сохранить журнал в текстовый файл")
        self.btn_log_save.setFixedHeight(self._px(32))
        self.btn_log_save.clicked.connect(self._save_log)
        log_toolbar.addWidget(self.btn_log_save)
        self.btn_log_clear = QPushButton("Очистить журнал")
        self.btn_log_clear.setToolTip("Очистить только журнал, не сбрасывая настройки")
        self.btn_log_clear.setFixedHeight(self._px(32))
        self.btn_log_clear.clicked.connect(self._clear_log)
        log_toolbar.addWidget(self.btn_log_clear)
        log_toolbar.addStretch()
        log_layout.addLayout(log_toolbar)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setFont(self._font(11, fixed=True))
        self.log_text.setMinimumHeight(self._px(160))
        log_layout.addWidget(self.log_text, 1)
        tabs.addTab(log_tab, "Журнал обработки")
        self.tabs = tabs
        self._apply_language()

        # Статус-бар: краткие подсказки и состояние
        self.status_bar = self.statusBar()
        self.status_bar.showMessage(self._t("Готов к работе", "Ready to work"))

        self._ensure_llm_settings_dialog()
        self._apply_language()

        # Esc — отмена текущей обработки
        esc = QAction(self)
        esc.setShortcut(QKeySequence(Qt.Key.Key_Escape))
        esc.triggered.connect(self._cancel_processing)
        self.addAction(esc)

        self._apply_theme()
        self._restore_geometry()

    # ──────────────────────────────────────────────────────────────
    # Меню, статус, геометрия окна, журнал
    # ──────────────────────────────────────────────────────────────

    def _build_menu_bar(self):
        menubar = self.menuBar()
        menubar.clear()

        self._menu_file = menubar.addMenu("Файл")
        file_menu = self._menu_file
        self._act_files = QAction("Выбрать файлы…", self)
        act_files = self._act_files
        act_files.setShortcut(QKeySequence.StandardKey.Open)
        act_files.setStatusTip("Добавить аудио- или видеофайлы в очередь")
        act_files.triggered.connect(self._select_files)
        file_menu.addAction(act_files)

        self._act_folder = QAction("Выбрать папку с файлами…", self)
        act_folder = self._act_folder
        act_folder.setStatusTip("Добавить все медиафайлы из папки и подпапок")
        act_folder.triggered.connect(self._select_files_folder)
        file_menu.addAction(act_folder)

        self._act_out = QAction("Папка сохранения…", self)
        act_out = self._act_out
        act_out.setStatusTip("Выбрать папку для результатов транскрибации")
        act_out.triggered.connect(self._select_output_folder)
        file_menu.addAction(act_out)

        file_menu.addSeparator()
        self._act_open_res = QAction("Открыть папку с результатами", self)
        act_open_res = self._act_open_res
        act_open_res.setStatusTip("Открыть папку с готовыми файлами")
        act_open_res.triggered.connect(self._open_results_folder)
        file_menu.addAction(act_open_res)

        file_menu.addSeparator()
        self._act_quit = QAction("Выход", self)
        act_quit = self._act_quit
        act_quit.setShortcut(QKeySequence.StandardKey.Quit)
        act_quit.triggered.connect(self.close)
        file_menu.addAction(act_quit)

        self._menu_view = menubar.addMenu("Вид")
        view_menu = self._menu_view
        self._act_theme = QAction("Переключить тему", self)
        self._act_theme.setShortcut(QKeySequence("Ctrl+T"))
        self._act_theme.setStatusTip("Светлая / тёмная тема оформления")
        self._act_theme.triggered.connect(self._toggle_theme)
        view_menu.addAction(self._act_theme)

        self._menu_settings = menubar.addMenu("Настройки")
        settings_menu = self._menu_settings
        self._act_device = QAction("Устройство (CPU / GPU)…", self)
        act_device = self._act_device
        act_device.setStatusTip("Выбрать CPU или видеокарту NVIDIA для распознавания")
        act_device.triggered.connect(self._change_device)
        settings_menu.addAction(act_device)

        settings_menu.addSeparator()
        self._act_llm = QAction("LLM API…", self)
        act_llm = self._act_llm
        act_llm.setStatusTip("Настроить API URL, ключ, модель и папку результатов LLM")
        act_llm.triggered.connect(self._open_llm_settings_dialog)
        settings_menu.addAction(act_llm)

        self._menu_help = menubar.addMenu("Справка")
        help_menu = self._menu_help
        self._act_about = QAction("О программе", self)
        act_about = self._act_about
        act_about.triggered.connect(self._show_about)
        help_menu.addAction(act_about)

    @staticmethod
    def _is_headless() -> bool:
        app = QApplication.instance()
        return bool(app) and app.platformName() in ("offscreen", "minimal")

    def _restore_geometry(self):
        # Геометрию из безголовых/тестовых сессий не восстанавливаем
        if self._is_headless():
            return
        saved = self.user_settings.settings.get("window_geometry")
        if saved:
            try:
                self.restoreGeometry(QByteArray.fromHex(saved.encode("ascii")))
            except Exception:
                pass

    def _save_geometry(self):
        if self._is_headless():
            return
        try:
            self.user_settings.settings["window_geometry"] = bytes(
                self.saveGeometry().toHex()
            ).decode("ascii")
            self.user_settings._save_settings()
        except Exception:
            pass

    def _translate_runtime_text(self, message: str) -> str:
        if self._lang == "ru" or not message:
            return message
        translated = str(message)
        replacements = [
            ("Подробности — на вкладке «Журнал обработки».", "See details in the 'Processing log' tab."),
            ("Не удалось сохранить журнал:\n", "Failed to save the log:\n"),
            ("Журнал скопирован в буфер обмена", "Log copied to clipboard"),
            ("Журнал сохранён: ", "Log saved: "),
            ("Журнал сохранён в ", "Log saved to "),
            ("Журнал очищен", "Log cleared"),
            ("Диаризация спикеров: ВКЛЮЧЕНА", "Speaker diarization: ENABLED"),
            ("Диаризация спикеров: ВЫКЛЮЧЕНА", "Speaker diarization: DISABLED"),
            ("Токен сохранён: ", "Token saved: "),
            ("Количество спикеров: автоопределение", "Speaker count: auto-detect"),
            ("Количество спикеров: ", "Speaker count: "),
            ("Обработка отменена пользователем", "Processing cancelled by user"),
            ("=== ОБРАБОТКА ЗАВЕРШЕНА ===", "=== PROCESSING FINISHED ==="),
            ("Общее время обработки: ", "Total processing time: "),
            (", с ошибками: ", ", with errors: "),
            ("Успешно: ", "Successful: "),
            ("Отменено. Обработано ", "Cancelled. Processed "),
            ("Готово с ошибками: ", "Completed with errors: "),
            (" успешно за ", " successful in "),
            ("Завершено за ", "Completed in "),
            ("Не удалось: ", "Failed: "),
            (" и ещё ", " and "),
            ("Критическая ошибка: ", "Critical error: "),
            ("Ошибка: ", "Error: "),
            ("Не удалось загрузить модель", "Failed to load model"),
            ("Анализ файлов и оценка времени обработки...", "Analyzing files and estimating processing time..."),
            ("Ожидаемое время обработки: ~", "Estimated processing time: ~"),
            ("Обработка ", "Processing "),
            (" файлов…", " files…"),
            ("Файл ", "File "),
            ("Оценка: ~", "Estimate: ~"),
            ("Осталось: ~", "Remaining: ~"),
            ("Подготовка…", "Preparing…"),
            ("Конвертация…", "Converting…"),
            ("Распознавание речи…", "Speech recognition…"),
            ("Распознавание речи (GigaAM-v3)...", "Speech recognition (GigaAM-v3)..."),
            ("Транскрибация завершена. Получено сегментов: ", "Transcription finished. Segments received: "),
            ("Пример структуры сегмента: keys=", "Example segment structure: keys="),
            ("Применение диаризации спикеров...", "Applying speaker diarization..."),
            ("Диаризация завершена. Найдено спикеров: ", "Diarization finished. Speakers found: "),
            ("Найдено сегментов речи: ", "Speech segments found: "),
            ("Сохранено символов: ", "Characters saved: "),
            ("Сохранено: ", "Saved: "),
            ("Время обработки: ", "Processing time: "),
            ("Конверсия: ", "Conversion: "),
            ("Транскрибация: ", "Transcription: "),
            ("Длительность: ", "Duration: "),
            ("Оценка времени обработки: ~", "Estimated processing time: ~"),
            ("неизвестна", "unknown"),
            ("ошибка определения длительности", "duration detection error"),
            ("длительность неизвестна", "duration unknown"),
            ("Ошибка при обработке файла ", "Error while processing file "),
            ("ОШИБКА при транскрибации: ", "TRANSCRIPTION ERROR: "),
            ("ОШИБКА VAD: ", "VAD ERROR: "),
            ("ПРЕДУПРЕЖДЕНИЕ: ", "WARNING: "),
            ("Ошибка при обработке ", "Error while processing "),
            ("Возможные причины:", "Possible reasons:"),
            ("Проверьте токен HF_TOKEN в .env файле и убедитесь, что приняли условия доступа:", "Check the HF_TOKEN in the .env file and make sure you accepted the access terms:"),
            ("Проверьте токен HF_TOKEN в src/config.py и убедитесь, что приняли условия доступа:", "Check the HF_TOKEN in src/config.py and make sure you accepted the access terms:"),
            ("Диаризация требует токен HuggingFace.", "Diarization requires a HuggingFace token."),
            ("Установите токен через чекбокс 'Диаризация' в интерфейсе.", "Set the token via the 'Diarization' checkbox in the interface."),
            ("Продолжаем без диаризации...", "Continuing without diarization..."),
            ("Спикер №", "Speaker #"),
        ]
        for old, new in replacements:
            translated = translated.replace(old, new)
        return translated

    def _set_status(self, message: str):
        """Дублирует ключевые сообщения в системный статус-бар."""
        if hasattr(self, "status_bar") and self.status_bar is not None:
            self.status_bar.showMessage(self._translate_runtime_text(message))

    def _copy_log(self):
        clipboard = QApplication.clipboard()
        if clipboard is not None:
            clipboard.setText(self.log_text.toPlainText())
            self._set_status(self._t("Журнал скопирован в буфер обмена", "Log copied to clipboard"))

    def _save_log(self):
        if not self.log_text.toPlainText().strip():
            QMessageBox.information(self, self._t("Журнал пуст", "Empty log"), self._t("Журнал пока пуст — нечего сохранять.", "The log is empty — nothing to save."))
            return
        initial_dir = self.user_settings.get_last_output_dir() or self.output_dir or os.path.expanduser("~")
        path, _ = QFileDialog.getSaveFileName(
            self, self._t("Сохранить журнал", "Save log"),
            os.path.join(initial_dir, "transcription_log.txt"),
            self._t("Текстовые файлы (*.txt);;Все файлы (*.*)", "Text files (*.txt);;All files (*.*)")
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.log_text.toPlainText())
            self._set_status(self._t(f"Журнал сохранён: {os.path.basename(path)}", f"Log saved: {os.path.basename(path)}"))
            self.log(f"Журнал сохранён в {path}")
        except OSError as e:
            QMessageBox.warning(self, self._t("Ошибка", "Error"), self._t("Не удалось сохранить журнал:\n", "Failed to save the log:\n") + str(e))

    def _clear_log(self):
        self.log_text.clear()
        self._set_status(self._t("Журнал очищен", "Log cleared"))

    def _open_results_folder(self):
        target = self._last_result_dir or self.output_dir or self.input_dir
        if not target or not os.path.isdir(target):
            QMessageBox.information(
                self,
                self._t("Папка недоступна", "Folder unavailable"),
                self._t("Папка с результатами ещё не определена.\nЗапустите обработку или выберите папку сохранения.", "The results folder has not been determined yet.\nStart processing or choose an output folder."),
            )
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(target))

    def _show_about(self):
        QMessageBox.about(
            self,
            self._t("О программе", "About"),
            (
                f"<b>{APP_TITLE}</b><br><br>"
                "Локальная транскрибация аудио и видео на модели <b>GigaAM v3</b> с поддержкой диаризации спикеров.<br><br>"
                "Возможности: пакетная обработка, загрузка по ссылке, таймкоды, экспорт в TXT / Markdown / SRT / VTT.<br><br>"
                "Поддерживаемые форматы ввода: mp3, wav, m4a, aac, flac, ogg, mp4, avi, mov, mkv, webm, wma, 3gp."
            ) if self._lang == "ru" else (
                f"<b>{APP_TITLE}</b><br><br>"
                "Local audio and video transcription powered by <b>GigaAM v3</b> with speaker diarization support.<br><br>"
                "Features: batch processing, URL download, timecodes, export to TXT / Markdown / SRT / VTT.<br><br>"
                "Supported input formats: mp3, wav, m4a, aac, flac, ogg, mp4, avi, mov, mkv, webm, wma, 3gp."
            )
        )

    _ACCENT_LIGHT = "#3b82f6"
    _CONVERSION_BAND = 0.15

    def _make_progress_bar(self, height: int, font_pt: int) -> QProgressBar:
        c = self._colors()
        bar = QProgressBar()
        scaled_height = self._px(height)
        bar.setFixedHeight(scaled_height)
        bar.setTextVisible(True)
        bar.setRange(0, 100)
        radius = scaled_height // 2
        r, r2 = c["progress_chunk"], c["progress_chunk2"]
        bar.setStyleSheet(
            f"QProgressBar {{ border: none; border-radius: {radius}px;"
            f"  background-color: {c['progress_bg']}; text-align: center; color: {c['text']};"
            f"  font-size: {self._pt_css(font_pt)}pt; font-weight: 600; }}"
            f"QProgressBar::chunk {{ border-radius: {radius}px;"
            f"  background-color: qlineargradient(x1:0,y1:0,x2:1,y2:0,"
            f"  stop:0 {r}, stop:1 {r2}); }}"
        )
        return bar

    def _create_progress_section(self, parent_layout):
        c = self._colors()
        progress_frame = QFrame()
        progress_frame.setObjectName("progress_card")
        frame_layout = QVBoxLayout(progress_frame)
        frame_layout.setContentsMargins(self._px(16), self._px(12), self._px(16), self._px(12))
        frame_layout.setSpacing(self._px(6))

        head_row = QHBoxLayout()
        self.lbl_overall = QLabel("Общий прогресс")
        lbl_overall = self.lbl_overall
        lbl_overall.setStyleSheet(self._transparent_label_style(c["text_sub"], font_pt=11, font_weight="bold"))
        head_row.addWidget(lbl_overall)
        head_row.addStretch()
        self.lbl_file_counter = QLabel("")
        self.lbl_file_counter.setStyleSheet(self._transparent_label_style(c["accent"], font_pt=11, font_weight="bold"))
        head_row.addWidget(self.lbl_file_counter)
        self.btn_cancel = QPushButton("Отменить")
        self.btn_cancel.setObjectName("cancel_button")
        self.btn_cancel.setToolTip("Остановить обработку после текущего файла  (Esc)")
        self.btn_cancel.setFixedHeight(self._px(28))
        self.btn_cancel.clicked.connect(self._cancel_processing)
        self.btn_cancel.setVisible(False)
        head_row.addWidget(self.btn_cancel)
        frame_layout.addLayout(head_row)

        self.progress_bar_total = self._make_progress_bar(height=22, font_pt=10)
        frame_layout.addWidget(self.progress_bar_total)

        self.detail_row = QWidget()
        detail_layout = QHBoxLayout(self.detail_row)
        detail_layout.setContentsMargins(0, 0, 0, 0)
        detail_layout.setSpacing(self._px(10))
        self.lbl_stage = QLabel("")
        self.lbl_stage.setStyleSheet(self._transparent_label_style(c["text_sub"], font_pt=9, font_weight="600"))
        detail_layout.addWidget(self.lbl_stage)
        detail_layout.addStretch()
        self.lbl_current_file = QLabel("")
        self.lbl_current_file.setStyleSheet(self._transparent_label_style(c["text_mute2"], font_pt=9))
        detail_layout.addWidget(self.lbl_current_file)
        frame_layout.addWidget(self.detail_row)
        self.detail_row.setVisible(False)

        self.progress_bar_file = self._make_progress_bar(height=16, font_pt=8)
        frame_layout.addWidget(self.progress_bar_file)

        self.lbl_status = QLabel(self._t("Готов к работе", "Ready to work"))
        self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_status.setFixedHeight(self._px(28))
        self.lbl_status.setStyleSheet(
            f"color: {c['status_text']}; font-size: {self._pt_css(10)}pt; font-weight: bold;"
            f"background-color: {c['status_bg']}; border-radius: {self._px(6)}px; padding: {self._px(2)}px;"
        )
        frame_layout.addWidget(self.lbl_status)

        self.btn_open_result = QPushButton("Открыть папку с результатами")
        self.btn_open_result.setObjectName("open_result_button")
        self.btn_open_result.setToolTip("Открыть папку с готовыми файлами в проводнике")
        self.btn_open_result.setFixedHeight(self._px(34))
        self.btn_open_result.clicked.connect(self._open_results_folder)
        self.btn_open_result.setVisible(False)
        frame_layout.addWidget(self.btn_open_result)

        parent_layout.addWidget(progress_frame)

    def _create_files_group(self) -> QGroupBox:
        group = QGroupBox("1. Выбор файлов")
        self.grp_files = group
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        main_layout.setSpacing(self._px(6))

        row1 = QHBoxLayout()
        row1.setSpacing(self._px(10))
        self.btn_select_files = QPushButton("Выбрать файлы")
        btn_select_files = self.btn_select_files
        btn_select_files.setToolTip("Выбрать аудио/видео файлы для обработки  (Ctrl+O)")
        btn_select_files.clicked.connect(self._select_files)
        btn_select_files.setFixedHeight(self._px(36))
        btn_select_files.setMinimumWidth(self._px(160))
        row1.addWidget(btn_select_files)

        self.btn_select_folder = QPushButton("Выбрать папку")
        btn_select_folder = self.btn_select_folder
        btn_select_folder.setToolTip("Добавить все медиафайлы из папки и подпапок")
        btn_select_folder.clicked.connect(self._select_files_folder)
        btn_select_folder.setFixedHeight(self._px(36))
        btn_select_folder.setMinimumWidth(self._px(150))
        row1.addWidget(btn_select_folder)

        self.input_path = QLineEdit()
        self.input_path.setPlaceholderText("Ссылка на медиа (YouTube и др.)")
        self.input_path.setToolTip("Вставьте ссылку и нажмите «Загрузить»")
        self.input_path.setFixedHeight(self._px(36))
        self.input_path.setMinimumWidth(self._px(200))
        self.input_path.returnPressed.connect(self._start_download)
        row1.addWidget(self.input_path, 1)

        self.btn_upload = QPushButton("Загрузить")
        self.btn_upload.setToolTip("Скачать медиа по ссылке и добавить в очередь")
        self.btn_upload.setFixedHeight(self._px(36))
        self.btn_upload.setMinimumWidth(self._px(100))
        self.btn_upload.clicked.connect(self._start_download)
        row1.addWidget(self.btn_upload)

        self.progress_upload = QProgressBar()
        self.progress_upload.setFixedHeight(self._px(36))
        self.progress_upload.setFixedWidth(self._px(90))
        self.progress_upload.setValue(0)
        self.progress_upload.setTextVisible(True)
        self.progress_upload.setVisible(False)
        row1.addWidget(self.progress_upload)
        main_layout.addLayout(row1)

        # Подсказка о выбранной папке источника
        self.lbl_input_folder = QLabel("Папка не выбрана")
        self.lbl_input_folder.setStyleSheet(self._transparent_label_style(self._colors()["text_mute"], font_pt=9))
        main_layout.addWidget(self.lbl_input_folder)

        # Очередь файлов + пустое состояние (drop-зона)
        self.drop_hint = QLabel("Перетащите сюда файлы или папки  ·  либо нажмите «Выбрать файлы»")
        self.drop_hint.setObjectName("drop_hint")
        self.drop_hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.drop_hint.setMinimumHeight(self._px(50))
        main_layout.addWidget(self.drop_hint)

        self.files_list = QListWidget()
        self.files_list.setObjectName("files_list")
        self.files_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.files_list.setMinimumHeight(self._px(72))
        self.files_list.setMaximumHeight(self._px(150))
        self.files_list.setToolTip("Очередь файлов. Выделите и нажмите Delete, чтобы убрать.")
        self.files_list.itemSelectionChanged.connect(self._update_files_controls)
        self.files_list.setVisible(False)
        main_layout.addWidget(self.files_list)

        controls = QHBoxLayout()
        controls.setSpacing(self._px(10))
        self.lbl_files_count = QLabel("Файлы не выбраны")
        self.lbl_files_count.setStyleSheet(self._transparent_label_style(self._colors()["text_mute"]))
        controls.addWidget(self.lbl_files_count)
        controls.addStretch()
        self.btn_remove_file = QPushButton("Убрать выбранное")
        self.btn_remove_file.setToolTip("Убрать выделенные файлы из очереди  (Delete)")
        self.btn_remove_file.setFixedHeight(self._px(32))
        self.btn_remove_file.setEnabled(False)
        self.btn_remove_file.clicked.connect(self._remove_selected_files)
        controls.addWidget(self.btn_remove_file)
        self.btn_clear_files = QPushButton("Очистить список")
        self.btn_clear_files.setToolTip("Убрать все файлы из очереди (настройки сохранятся)")
        self.btn_clear_files.setFixedHeight(self._px(32))
        self.btn_clear_files.setEnabled(False)
        self.btn_clear_files.clicked.connect(self._clear_files_list)
        controls.addWidget(self.btn_clear_files)
        main_layout.addLayout(controls)

        group.setLayout(main_layout)
        return group

    def _size_list_to_contents(self, list_widget: QListWidget, min_rows: int = 1, max_rows: int = 5):
        """Выставляет высоту списка по числу строк, вместо фиксированного
        большого блока с пустым местом для 1-2 файлов."""
        row_h = list_widget.sizeHintForRow(0)
        if row_h <= 0:
            row_h = self._px(24)
        count = max(min_rows, min(max_rows, list_widget.count() or min_rows))
        frame = self._px(2) * 2 + self._px(4)
        list_widget.setFixedHeight(row_h * count + frame)

    def _refresh_files_list(self):
        """Синхронизирует QListWidget и пустое состояние с self.files_to_process."""
        if not hasattr(self, "files_list"):
            return
        self.files_list.clear()
        for path in self.files_to_process:
            item = QListWidgetItem(os.path.basename(path))
            item.setToolTip(path)
            item.setData(Qt.ItemDataRole.UserRole, path)
            self.files_list.addItem(item)
        has_files = bool(self.files_to_process)
        self.files_list.setVisible(has_files)
        self.drop_hint.setVisible(not has_files)
        if has_files:
            self._size_list_to_contents(self.files_list)
        c = self._colors()
        if has_files:
            self.lbl_files_count.setText(self._t(f"Выбрано файлов: {len(self.files_to_process)}", f"Selected files: {len(self.files_to_process)}"))
            self.lbl_files_count.setStyleSheet(self._transparent_label_style(c["text_sub"]))
        else:
            self.lbl_files_count.setText("Файлы не выбраны")
            self.lbl_files_count.setStyleSheet(self._transparent_label_style(c["text_mute"]))
        self._update_files_controls()

    def _update_files_controls(self):
        if not hasattr(self, "btn_clear_files"):
            return
        has_files = bool(self.files_to_process)
        self.btn_clear_files.setEnabled(has_files and not self.is_processing)
        self.btn_remove_file.setEnabled(
            bool(self.files_list.selectedItems()) and not self.is_processing
        )

    def _remove_selected_files(self):
        if self.is_processing:
            return
        selected = {
            item.data(Qt.ItemDataRole.UserRole) for item in self.files_list.selectedItems()
        }
        if not selected:
            return
        removed = len(selected)
        self.files_to_process = [p for p in self.files_to_process if p not in selected]
        self._refresh_files_list()
        self.log(f"Убрано из очереди: {removed} файлов")

    def _clear_files_list(self):
        if self.is_processing or not self.files_to_process:
            return
        self.files_to_process = []
        self._refresh_files_list()
        self.log("Очередь файлов очищена")

    def keyPressEvent(self, event):
        if (
            event.key() == Qt.Key.Key_Delete
            and hasattr(self, "files_list")
            and self.files_list.hasFocus()
        ):
            self._remove_selected_files()
            return
        if (
            event.key() == Qt.Key.Key_Delete
            and hasattr(self, "llm_files_list")
            and self.llm_files_list.hasFocus()
        ):
            self._remove_selected_llm_files()
            return
        super().keyPressEvent(event)

    def _style_drop_hint(self):
        c = self._colors()
        active = getattr(self, "_drop_active", False)
        border = c["accent"] if active else c["border"]
        bg = c["btn_hover_bg"] if active else c["bg_card"]
        text = c["accent"] if active else c["text_mute2"]
        if hasattr(self, "drop_hint"):
            self.drop_hint.setStyleSheet(
                f"#drop_hint {{ border: 2px dashed {border}; border-radius: {self._px(10)}px;"
                f"  background-color: {bg}; color: {text};"
                f"  font-size: {self._pt_css(11)}pt; padding: {self._px(8)}px; }}"
            )
        if hasattr(self, "llm_drop_hint"):
            self.llm_drop_hint.setStyleSheet(
                f"#llm_drop_hint {{ border: 2px dashed {border}; border-radius: {self._px(10)}px;"
                f"  background-color: {bg}; color: {text};"
                f"  font-size: {self._pt_css(11)}pt; padding: {self._px(8)}px; }}"
            )

    def _create_output_group(self) -> QGroupBox:
        group = QGroupBox("2. Папка сохранения результатов")
        self.grp_output = group
        layout = QHBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(10))
        layout.setSpacing(self._px(12))
        self.btn_output_select = QPushButton("Выбрать папку")
        btn_output = self.btn_output_select
        btn_output.clicked.connect(self._select_output_folder)
        btn_output.setMinimumWidth(self._px(220))
        btn_output.setFixedHeight(self._px(36))
        layout.addWidget(btn_output)
        self.lbl_output_folder = QLabel("Папка не выбрана (по умолчанию - рядом с файлом)")
        self.lbl_output_folder.setStyleSheet(self._transparent_label_style(self._colors()["text_mute"]))
        layout.addWidget(self.lbl_output_folder, 1)
        group.setLayout(layout)
        return group

    def _create_diarization_group(self) -> QGroupBox:
        group = QGroupBox("3. Диаризация спикеров")
        self.grp_diarization = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(10))
        layout.setSpacing(self._px(8))
        self.cb_diarization = QCheckBox("Включить диаризацию спикеров")
        self.cb_diarization.setToolTip("Определять, кто из спикеров говорит (нужен HF_TOKEN)")
        self.cb_diarization.stateChanged.connect(self._toggle_diarization)
        layout.addWidget(self.cb_diarization)
        speakers_layout = QHBoxLayout()
        speakers_layout.setSpacing(self._px(12))
        self.lbl_num_speakers = QLabel("Кол-во спикеров:")
        speakers_layout.addWidget(self.lbl_num_speakers)
        self.entry_num_speakers = QSpinBox()
        self.entry_num_speakers.setRange(0, 20)
        self.entry_num_speakers.setValue(0)
        self.entry_num_speakers.setSpecialValueText("Авто")
        self.entry_num_speakers.setToolTip("0 = автоопределение количества спикеров")
        self.entry_num_speakers.setEnabled(False)
        self.entry_num_speakers.setFixedHeight(self._px(32))
        self.entry_num_speakers.setMinimumWidth(self._px(140))
        self.entry_num_speakers.setMaximumWidth(self._px(200))
        speakers_layout.addWidget(self.entry_num_speakers)
        speakers_layout.addStretch()
        layout.addLayout(speakers_layout)
        self.lbl_diarization_info = QLabel("Автоматическое определение спикеров (требуется HF_TOKEN)")
        info_label = self.lbl_diarization_info
        info_label.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(info_label)
        group.setLayout(layout)
        return group

    def _create_formats_group(self) -> QGroupBox:
        group = QGroupBox("4. Форматы вывода")
        self.grp_formats = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(10))
        layout.setSpacing(self._px(6))
        self.format_checkboxes = {}

        row1 = QHBoxLayout()
        row1.setSpacing(self._px(20))
        for fmt in ['txt', 'txt_timecodes', 'txt_diarize', 'txt_diarize_timecodes']:
            cb = QCheckBox(OUTPUT_FORMATS[fmt])
            cb.setChecked(fmt in ('txt', 'txt_timecodes'))
            if fmt in ('txt_diarize', 'txt_diarize_timecodes'):
                cb.setEnabled(False)
            cb.stateChanged.connect(lambda state, f=fmt: self._toggle_format(f))
            row1.addWidget(cb)
            self.format_checkboxes[fmt] = cb
        row1.addStretch()
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.setSpacing(self._px(20))
        for fmt in ('md', 'srt', 'vtt'):
            cb = QCheckBox(OUTPUT_FORMATS[fmt])
            cb.setChecked(False)
            cb.stateChanged.connect(lambda state, f=fmt: self._toggle_format(f))
            row2.addWidget(cb)
            self.format_checkboxes[fmt] = cb
        row2.addStretch()
        layout.addLayout(row2)

        group.setLayout(layout)
        return group

    def _create_llm_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(self._px(8), self._px(14), self._px(8), self._px(6))
        layout.setSpacing(self._px(4))

        layout.addWidget(self._create_llm_source_group())
        layout.addWidget(self._create_llm_output_group())
        layout.addWidget(self._create_llm_actions_group())
        layout.addWidget(self._create_llm_save_group())

        self.btn_llm_process = QPushButton("ОБРАБОТАТЬ")
        self.btn_llm_process.setObjectName("start_button")
        self.btn_llm_process.setFixedHeight(self._px(52))
        self.btn_llm_process.setToolTip("Запустить LLM-обработку выбранных транскриптов")
        self.btn_llm_process.clicked.connect(self._start_llm_processing)
        layout.addWidget(self.btn_llm_process)

        layout.addWidget(self._create_llm_result_group(), 1)

        self.btn_llm_clear = QPushButton("ОЧИСТИТЬ ВСЕ")
        self.btn_llm_clear.setObjectName("clear_button")
        self.btn_llm_clear.setFixedHeight(self._px(40))
        self.btn_llm_clear.setToolTip("Сбросить выбранные транскрипты, ручной текст и результат LLM")
        self.btn_llm_clear.clicked.connect(self._clear_llm_all)
        layout.addWidget(self.btn_llm_clear)

        return tab

    def _create_llm_api_group(self) -> QGroupBox:
        group = QGroupBox("LLM API")
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(10))
        layout.setSpacing(self._px(8))

        label_col = self._label_column_width(("Провайдер:", "Модель:", "API URL:", "API Key:"))
        self.llm_provider_labels = {}
        self.llm_provider_items = {5: "Другое"}

        provider_row = QHBoxLayout()
        self.llm_provider_labels["provider"] = self._form_label("Провайдер:", label_col)
        provider_row.addWidget(self.llm_provider_labels["provider"])
        self.combo_llm_provider = QComboBox()
        self.combo_llm_provider.addItems(["API", "Claude Code", "Codex", "OpenCode", "Pi", "Другое"])
        self.combo_llm_provider.setMinimumWidth(self._px(180))
        provider_row.addWidget(self.combo_llm_provider)
        provider_row.addStretch()
        layout.addLayout(provider_row)

        common_row = QHBoxLayout()
        self.llm_provider_labels["model"] = self._form_label("Модель:", label_col)
        common_row.addWidget(self.llm_provider_labels["model"])
        self.entry_llm_model = QLineEdit()
        self.entry_llm_model.setPlaceholderText("gpt-4.1-mini / sonnet / o3 / qwen ...")
        common_row.addWidget(self.entry_llm_model, 1)
        common_row.addSpacing(self._px(12))
        common_row.addWidget(QLabel("Temperature:"))
        self.entry_llm_temperature = QLineEdit()
        self.entry_llm_temperature.setMaximumWidth(self._px(110))
        common_row.addWidget(self.entry_llm_temperature)
        layout.addLayout(common_row)

        self.llm_api_settings_widget = QWidget()
        self.llm_api_settings_widget.setStyleSheet("background: transparent;")
        api_layout = QVBoxLayout(self.llm_api_settings_widget)
        api_layout.setContentsMargins(0, 0, 0, 0)
        api_layout.setSpacing(self._px(8))
        row1 = QHBoxLayout()
        row1.addWidget(self._form_label("API URL:", label_col))
        self.entry_llm_api_url = QLineEdit()
        self.entry_llm_api_url.setPlaceholderText("https://api.openai.com/v1")
        row1.addWidget(self.entry_llm_api_url, 1)
        api_layout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(self._form_label("API Key:", label_col))
        self.entry_llm_api_key = QLineEdit()
        self.entry_llm_api_key.setEchoMode(QLineEdit.EchoMode.Password)
        self.entry_llm_api_key.setPlaceholderText("Bearer token / API key")
        row2.addWidget(self.entry_llm_api_key, 1)
        api_layout.addLayout(row2)
        layout.addWidget(self.llm_api_settings_widget)

        self.llm_claude_settings_widget = QWidget()
        self.llm_claude_settings_widget.setStyleSheet("background: transparent;")
        claude_layout = QVBoxLayout(self.llm_claude_settings_widget)
        claude_layout.setContentsMargins(0, 0, 0, 0)
        claude_layout.setSpacing(self._px(8))
        claude_col = self._label_column_width(("Claude Code путь:", "Claude доп. аргументы:"))
        row4 = QHBoxLayout()
        self.llm_provider_labels["claude_path"] = self._form_label("Claude Code путь:", claude_col)
        row4.addWidget(self.llm_provider_labels["claude_path"])
        self.entry_llm_claude_path = QLineEdit()
        self.entry_llm_claude_path.setPlaceholderText("claude")
        row4.addWidget(self.entry_llm_claude_path, 1)
        claude_layout.addLayout(row4)
        row5 = QHBoxLayout()
        self.llm_provider_labels["claude_args"] = self._form_label("Claude доп. аргументы:", claude_col)
        row5.addWidget(self.llm_provider_labels["claude_args"])
        self.entry_llm_claude_args = QLineEdit()
        self.entry_llm_claude_args.setPlaceholderText("например: --permission-mode bypassPermissions")
        row5.addWidget(self.entry_llm_claude_args, 1)
        claude_layout.addLayout(row5)
        layout.addWidget(self.llm_claude_settings_widget)

        self.llm_codex_settings_widget = QWidget()
        self.llm_codex_settings_widget.setStyleSheet("background: transparent;")
        codex_layout = QVBoxLayout(self.llm_codex_settings_widget)
        codex_layout.setContentsMargins(0, 0, 0, 0)
        codex_layout.setSpacing(self._px(8))
        codex_col = self._label_column_width(("Codex путь:", "Codex доп. аргументы:"))
        row6 = QHBoxLayout()
        self.llm_provider_labels["codex_path"] = self._form_label("Codex путь:", codex_col)
        row6.addWidget(self.llm_provider_labels["codex_path"])
        self.entry_llm_codex_path = QLineEdit()
        self.entry_llm_codex_path.setPlaceholderText("codex")
        row6.addWidget(self.entry_llm_codex_path, 1)
        codex_layout.addLayout(row6)
        row7 = QHBoxLayout()
        self.llm_provider_labels["codex_args"] = self._form_label("Codex доп. аргументы:", codex_col)
        row7.addWidget(self.llm_provider_labels["codex_args"])
        self.entry_llm_codex_args = QLineEdit()
        self.entry_llm_codex_args.setPlaceholderText("например: --dangerously-bypass-approvals-and-sandbox")
        row7.addWidget(self.entry_llm_codex_args, 1)
        codex_layout.addLayout(row7)
        layout.addWidget(self.llm_codex_settings_widget)

        self.llm_opencode_settings_widget = QWidget()
        self.llm_opencode_settings_widget.setStyleSheet("background: transparent;")
        opencode_layout = QVBoxLayout(self.llm_opencode_settings_widget)
        opencode_layout.setContentsMargins(0, 0, 0, 0)
        opencode_layout.setSpacing(self._px(8))
        opencode_col = self._label_column_width(("OpenCode путь:", "OpenCode доп. аргументы:"))
        row8 = QHBoxLayout()
        self.llm_provider_labels["opencode_path"] = self._form_label("OpenCode путь:", opencode_col)
        row8.addWidget(self.llm_provider_labels["opencode_path"])
        self.entry_llm_opencode_path = QLineEdit()
        self.entry_llm_opencode_path.setPlaceholderText("opencode")
        row8.addWidget(self.entry_llm_opencode_path, 1)
        opencode_layout.addLayout(row8)
        row9 = QHBoxLayout()
        self.llm_provider_labels["opencode_args"] = self._form_label("OpenCode доп. аргументы:", opencode_col)
        row9.addWidget(self.llm_provider_labels["opencode_args"])
        self.entry_llm_opencode_args = QLineEdit()
        self.entry_llm_opencode_args.setPlaceholderText("например: --print")
        row9.addWidget(self.entry_llm_opencode_args, 1)
        opencode_layout.addLayout(row9)
        layout.addWidget(self.llm_opencode_settings_widget)

        self.llm_pi_settings_widget = QWidget()
        self.llm_pi_settings_widget.setStyleSheet("background: transparent;")
        pi_layout = QVBoxLayout(self.llm_pi_settings_widget)
        pi_layout.setContentsMargins(0, 0, 0, 0)
        pi_layout.setSpacing(self._px(8))
        pi_col = self._label_column_width(("Pi путь:", "Pi provider:", "Pi доп. аргументы:"))
        row10 = QHBoxLayout()
        self.llm_provider_labels["pi_path"] = self._form_label("Pi путь:", pi_col)
        row10.addWidget(self.llm_provider_labels["pi_path"])
        self.entry_llm_pi_path = QLineEdit()
        self.entry_llm_pi_path.setPlaceholderText("pi")
        row10.addWidget(self.entry_llm_pi_path, 1)
        pi_layout.addLayout(row10)
        row11 = QHBoxLayout()
        self.llm_provider_labels["pi_provider"] = self._form_label("Pi provider:", pi_col)
        row11.addWidget(self.llm_provider_labels["pi_provider"])
        self.entry_llm_pi_provider = QLineEdit()
        self.entry_llm_pi_provider.setPlaceholderText("openai / anthropic / google ...")
        row11.addWidget(self.entry_llm_pi_provider, 1)
        pi_layout.addLayout(row11)
        row12 = QHBoxLayout()
        self.llm_provider_labels["pi_args"] = self._form_label("Pi доп. аргументы:", pi_col)
        row12.addWidget(self.llm_provider_labels["pi_args"])
        self.entry_llm_pi_args = QLineEdit()
        self.entry_llm_pi_args.setPlaceholderText("например: --no-tools --thinking low")
        row12.addWidget(self.entry_llm_pi_args, 1)
        pi_layout.addLayout(row12)
        layout.addWidget(self.llm_pi_settings_widget)

        self.llm_other_settings_widget = QWidget()
        self.llm_other_settings_widget.setStyleSheet("background: transparent;")
        other_layout = QVBoxLayout(self.llm_other_settings_widget)
        other_layout.setContentsMargins(0, 0, 0, 0)
        other_layout.setSpacing(self._px(8))
        other_col = self._label_column_width(("Команда:", "Аргументы:"))
        row13 = QHBoxLayout()
        self.llm_provider_labels["other_path"] = self._form_label("Команда:", other_col)
        row13.addWidget(self.llm_provider_labels["other_path"])
        self.entry_llm_other_path = QLineEdit()
        self.entry_llm_other_path.setPlaceholderText("путь к CLI, например my-llm")
        row13.addWidget(self.entry_llm_other_path, 1)
        other_layout.addLayout(row13)
        row14 = QHBoxLayout()
        self.llm_provider_labels["other_args"] = self._form_label("Аргументы:", other_col)
        row14.addWidget(self.llm_provider_labels["other_args"])
        self.entry_llm_other_args = QLineEdit()
        self.entry_llm_other_args.setPlaceholderText("аргументы; промпт будет добавлен в конец как последний параметр")
        row14.addWidget(self.entry_llm_other_args, 1)
        other_layout.addLayout(row14)
        layout.addWidget(self.llm_other_settings_widget)

        self.lbl_llm_provider_info = QLabel()
        self.lbl_llm_provider_info.setWordWrap(True)
        self.lbl_llm_provider_info.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(self.lbl_llm_provider_info)

        self.combo_llm_provider.currentTextChanged.connect(self._update_llm_provider_fields)
        self._update_llm_provider_fields(self.combo_llm_provider.currentText())
        group.setLayout(layout)
        return group

    def _update_llm_provider_fields(self, provider: str):
        provider = self._normalize_llm_provider(provider)
        widgets = {
            "API": self.llm_api_settings_widget,
            "Claude Code": self.llm_claude_settings_widget,
            "Codex": self.llm_codex_settings_widget,
            "OpenCode": self.llm_opencode_settings_widget,
            "Pi": self.llm_pi_settings_widget,
            "Other": self.llm_other_settings_widget,
        }
        for name, widget in widgets.items():
            widget.setVisible(name == provider)

        info_map = {
            "API": self._t("Режим автоопределения API: поддерживает OpenAI-compatible и Anthropic Messages API. localhost/local network тоже поддерживается, если сервер совместим с одним из этих форматов.", "API auto-detection mode: supports OpenAI-compatible APIs and the Anthropic Messages API. localhost/local network is also supported if the server is compatible with one of these formats."),
            "Claude Code": self._t("Локальный Claude CLI. Используются путь к claude, модель и доп. аргументы.", "Local Claude CLI. Uses the claude path, model, and extra arguments."),
            "Codex": self._t("Локальный Codex CLI. Используются путь к codex, модель и доп. аргументы.", "Local Codex CLI. Uses the codex path, model, and extra arguments."),
            "OpenCode": self._t("Локальный OpenCode CLI. Будет запущен как команда + аргументы + промпт в конце.", "Local OpenCode CLI. It will be launched as command + arguments + prompt at the end."),
            "Pi": self._t("Локальный pi CLI. Можно указать внутренний provider для pi, модель и доп. аргументы.", "Local pi CLI. You can specify the internal provider for pi, the model, and extra arguments."),
            "Other": self._t("Произвольный CLI. Укажи команду и аргументы; промпт будет передан последним аргументом.", "Arbitrary CLI. Specify the command and arguments; the prompt will be passed as the last argument."),
        }
        self.lbl_llm_provider_info.setText(info_map.get(provider, ""))

    def _ensure_llm_settings_dialog(self):
        if getattr(self, "_llm_settings_dialog", None) is not None:
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("Настройки LLM")
        dialog.setMinimumWidth(self._px(760))
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(self._px(12), self._px(12), self._px(12), self._px(12))
        layout.setSpacing(self._px(8))
        self.grp_llm_api_settings = self._create_llm_api_group()
        layout.addWidget(self.grp_llm_api_settings)

        self.prompts_group = QGroupBox("Готовые промпты")
        prompts_layout = QVBoxLayout()
        prompts_layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        prompts_layout.setSpacing(self._px(8))

        self.lbl_llm_summary_prompt = QLabel("Промпт для выжимки:")
        prompts_layout.addWidget(self.lbl_llm_summary_prompt)
        self.txt_llm_summary_prompt = QTextEdit()
        self.txt_llm_summary_prompt.setMinimumHeight(self._px(100))
        prompts_layout.addWidget(self.txt_llm_summary_prompt)

        self.lbl_llm_tasks_prompt = QLabel("Промпт для задач:")
        prompts_layout.addWidget(self.lbl_llm_tasks_prompt)
        self.txt_llm_tasks_prompt = QTextEdit()
        self.txt_llm_tasks_prompt.setMinimumHeight(self._px(100))
        prompts_layout.addWidget(self.txt_llm_tasks_prompt)
        self.txt_llm_custom_prompt = QTextEdit()
        self.prompts_group.setLayout(prompts_layout)
        layout.addWidget(self.prompts_group)

        self.lbl_llm_settings_note = QLabel("Можно использовать OpenAI-compatible API, Anthropic Messages API, а также локальные Claude Code / Codex / OpenCode / Pi. Для API режим сам определяет тип API по URL или endpoint. Выбранный провайдер, модель, temperature, чекбоксы, prompt и файлы сохраняются между запусками. API Key лучше хранить в .env.")
        self.lbl_llm_settings_note.setWordWrap(True)
        self.lbl_llm_settings_note.setContentsMargins(self._px(4), self._px(2), self._px(4), self._px(2))
        self.lbl_llm_settings_note.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(self.lbl_llm_settings_note)

        self._llm_settings_buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Save | QDialogButtonBox.StandardButton.Close)
        self._llm_settings_buttons.button(QDialogButtonBox.StandardButton.Save).setText("Сохранить")
        self._llm_settings_buttons.button(QDialogButtonBox.StandardButton.Close).setText("Закрыть")
        self._llm_settings_buttons.accepted.connect(self._save_llm_settings_from_dialog)
        self._llm_settings_buttons.rejected.connect(dialog.reject)
        self._llm_settings_buttons.button(QDialogButtonBox.StandardButton.Close).clicked.connect(dialog.accept)
        layout.addWidget(self._llm_settings_buttons)
        self._llm_settings_dialog = dialog

    def _save_llm_settings_from_dialog(self):
        try:
            self._collect_llm_settings()
        except ValueError as e:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), str(e))
            return
        self._save_ui_settings()
        QMessageBox.information(self, self._t("Настройки", "Settings"), self._t("LLM-настройки сохранены", "LLM settings saved"))

    def _open_llm_settings_dialog(self):
        self._ensure_llm_settings_dialog()
        self._llm_settings_dialog.exec()

    def _create_llm_actions_group(self) -> QGroupBox:
        group = QGroupBox("3. Что сделать")
        self.grp_llm_actions = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        layout.setSpacing(self._px(6))
        self.llm_action_checkboxes = {}

        row = QHBoxLayout()
        row.setSpacing(self._px(20))
        for key, label, checked in (
            ("summary", "Выжимка", True),
            ("tasks", "Задачи", False),
            ("custom", "Свой промпт", False),
        ):
            cb = QCheckBox(label)
            cb.setChecked(checked)
            row.addWidget(cb)
            self.llm_action_checkboxes[key] = cb
        row.addStretch()
        layout.addLayout(row)

        self.lbl_llm_actions_note = QLabel("Отметьте один или несколько режимов обработки. Для «Свой промпт» текст задается в меню «Настройки → LLM API…».")
        self.lbl_llm_actions_note.setWordWrap(True)
        self.lbl_llm_actions_note.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(self.lbl_llm_actions_note)
        group.setLayout(layout)
        return group

    def _create_llm_output_group(self) -> QGroupBox:
        group = QGroupBox("2. Куда сохранить")
        self.grp_llm_output = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        layout.setSpacing(self._px(6))

        row = QHBoxLayout()
        row.setSpacing(self._px(10))
        self.btn_llm_output = QPushButton("Выбрать папку")
        btn_output = self.btn_llm_output
        btn_output.setFixedHeight(self._px(36))
        btn_output.clicked.connect(self._select_llm_output_folder)
        row.addWidget(btn_output)
        self.lbl_llm_output = QLabel("Папка не выбрана (по умолчанию - рядом с транскриптом)")
        self.lbl_llm_output.setStyleSheet(self._transparent_label_style(self._colors()["text_mute"]))
        row.addWidget(self.lbl_llm_output, 1)
        layout.addLayout(row)

        self.lbl_llm_output_note = QLabel("Если папка не выбрана, результат будет сохранен рядом с исходным транскриптом.")
        self.lbl_llm_output_note.setWordWrap(True)
        self.lbl_llm_output_note.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(self.lbl_llm_output_note)
        group.setLayout(layout)
        return group

    def _create_llm_save_group(self) -> QGroupBox:
        group = QGroupBox("4. Форматы вывода")
        self.grp_llm_save = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        layout.setSpacing(self._px(6))
        self.llm_export_checkboxes = {}

        row1 = QHBoxLayout()
        row1.setSpacing(self._px(20))
        for key, label, checked in (("txt", "TXT (.txt)", True), ("md", "Markdown (.md)", False), ("docx", "DOCX (.docx)", False)):
            cb = QCheckBox(label)
            cb.setChecked(checked)
            row1.addWidget(cb)
            self.llm_export_checkboxes[key] = cb
        row1.addStretch()
        layout.addLayout(row1)
        group.setLayout(layout)
        return group

    def _create_llm_source_group(self) -> QGroupBox:
        group = QGroupBox("1. Источник транскрипта")
        self.grp_llm_source = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        layout.setSpacing(self._px(6))

        files_row = QHBoxLayout()
        self.btn_select_transcripts = QPushButton("Выбрать транскрипты")
        btn_select_transcripts = self.btn_select_transcripts
        btn_select_transcripts.setFixedHeight(self._px(36))
        btn_select_transcripts.clicked.connect(self._select_llm_transcript_files)
        files_row.addWidget(btn_select_transcripts)
        files_row.addStretch()
        layout.addLayout(files_row)
        # Совместимость с кодом, обращающимся к self.lbl_llm_files напрямую в lbl_llm_files_count
        # (единая видимая метка статуса вместо двух дублирующих надписей).

        self.llm_drop_hint = QLabel("Перетащите сюда транскрипты  ·  либо нажмите «Выбрать транскрипты»")
        self.llm_drop_hint.setObjectName("llm_drop_hint")
        self.llm_drop_hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.llm_drop_hint.setMinimumHeight(self._px(50))
        layout.addWidget(self.llm_drop_hint)

        self.llm_files_list = QListWidget()
        self.llm_files_list.setObjectName("llm_files_list")
        self.llm_files_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.llm_files_list.setMinimumHeight(self._px(72))
        self.llm_files_list.setMaximumHeight(self._px(150))
        self.llm_files_list.setToolTip("Список транскриптов. Выделите и нажмите Delete, чтобы убрать.")
        self.llm_files_list.itemSelectionChanged.connect(self._update_llm_files_controls)
        self.llm_files_list.setVisible(False)
        layout.addWidget(self.llm_files_list)

        controls = QHBoxLayout()
        controls.setSpacing(self._px(10))
        self.lbl_llm_files_count = QLabel("Файлы не выбраны")
        self.lbl_llm_files_count.setStyleSheet(self._transparent_label_style(self._colors()["text_mute"]))
        controls.addWidget(self.lbl_llm_files_count)
        self.lbl_llm_files = self.lbl_llm_files_count
        controls.addStretch()
        self.btn_remove_llm_file = QPushButton("Убрать выбранное")
        self.btn_remove_llm_file.setFixedHeight(self._px(32))
        self.btn_remove_llm_file.setEnabled(False)
        self.btn_remove_llm_file.clicked.connect(self._remove_selected_llm_files)
        controls.addWidget(self.btn_remove_llm_file)
        self.btn_clear_llm_files = QPushButton("Очистить список")
        self.btn_clear_llm_files.setFixedHeight(self._px(32))
        self.btn_clear_llm_files.setEnabled(False)
        self.btn_clear_llm_files.clicked.connect(self._clear_llm_files_list)
        controls.addWidget(self.btn_clear_llm_files)
        layout.addLayout(controls)

        self.lbl_llm_supported = QLabel("Поддерживаемые файлы: .txt, .md, .srt, .vtt — либо вставьте транскрипт вручную ниже")
        info = self.lbl_llm_supported
        info.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(info)

        self.txt_llm_transcript = QTextEdit()
        self.txt_llm_transcript.setPlaceholderText("Вставьте сюда транскрипт, если не хотите выбирать файлы")
        self.txt_llm_transcript.setMinimumHeight(self._px(170))
        layout.addWidget(self.txt_llm_transcript)
        group.setLayout(layout)
        return group

    def _create_llm_prompt_group(self) -> QGroupBox:
        group = QGroupBox("2. Промпт")
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(9))
        layout.setSpacing(self._px(6))

        hint = QLabel("Все промпты настраиваются в меню «Настройки → LLM API…». Здесь достаточно выбрать режимы обработки. Для режима «Свой промпт» заранее заполните пользовательский промпт в настройках.")
        hint.setWordWrap(True)
        hint.setStyleSheet(self._transparent_label_style(self._colors()["text_mute2"], font_pt=9))
        layout.addWidget(hint)
        group.setLayout(layout)
        return group

    def _create_llm_result_group(self) -> QGroupBox:
        group = QGroupBox("5. Результат LLM")
        self.grp_llm_result = group
        layout = QVBoxLayout()
        layout.setContentsMargins(self._px(12), self._px(8), self._px(12), self._px(8))
        layout.setSpacing(self._px(6))

        self.lbl_llm_status = QLabel("Готово к LLM-обработке")
        self.lbl_llm_status.setWordWrap(True)
        self.lbl_llm_status.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.lbl_llm_status.setStyleSheet(
            self._transparent_label_style(self._colors()["text_sub"], font_pt=10, font_weight="bold")
        )
        layout.addWidget(self.lbl_llm_status)

        self.txt_llm_result = QTextEdit()
        self.txt_llm_result.setReadOnly(True)
        self.txt_llm_result.setFont(self._font(10, fixed=True))
        self.txt_llm_result.setMinimumHeight(self._px(260))
        layout.addWidget(self.txt_llm_result, 1)
        group.setLayout(layout)
        return group

    # ──────────────────────────────────────────────────────────────
    # Диалог HF токена
    # ──────────────────────────────────────────────────────────────

    def _show_hf_token_dialog(self) -> bool:
        dlg = QDialog(self)
        dlg.setWindowTitle(self._t("HuggingFace токен для диаризации", "HuggingFace token for diarization"))
        dlg.setMinimumWidth(self._px(520))
        layout = QVBoxLayout(dlg)
        layout.setSpacing(self._px(10))
        info = QLabel(
            self._t(
                "<b>Диаризация спикеров требует HuggingFace токен</b><br><br>"
                "1. Создайте аккаунт на <a href='https://huggingface.co'>huggingface.co</a><br>"
                "2. Получите токен: <a href='https://huggingface.co/settings/tokens'>huggingface.co/settings/tokens</a><br>"
                "3. Примите условия доступа к моделям:<br>"
                "&nbsp;&nbsp;&nbsp;<a href='https://huggingface.co/pyannote/speaker-diarization-3.1'>pyannote/speaker-diarization-3.1</a><br>"
                "&nbsp;&nbsp;&nbsp;<a href='https://huggingface.co/pyannote/segmentation-3.0'>pyannote/segmentation-3.0</a><br><br>"
                "Вставьте ваш токен ниже (начинается с <b>hf_</b>):",
                "<b>Speaker diarization requires a HuggingFace token</b><br><br>"
                "1. Create an account at <a href='https://huggingface.co'>huggingface.co</a><br>"
                "2. Get a token: <a href='https://huggingface.co/settings/tokens'>huggingface.co/settings/tokens</a><br>"
                "3. Accept the access terms for the models:<br>"
                "&nbsp;&nbsp;&nbsp;<a href='https://huggingface.co/pyannote/speaker-diarization-3.1'>pyannote/speaker-diarization-3.1</a><br>"
                "&nbsp;&nbsp;&nbsp;<a href='https://huggingface.co/pyannote/segmentation-3.0'>pyannote/segmentation-3.0</a><br><br>"
                "Paste your token below (it starts with <b>hf_</b>):"
            )
        )
        info.setOpenExternalLinks(True)
        info.setWordWrap(True)
        layout.addWidget(info)
        token_input = QLineEdit()
        token_input.setPlaceholderText("hf_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")
        token_input.setEchoMode(QLineEdit.EchoMode.Password)
        current_token = os.getenv("HF_TOKEN", "")
        if current_token:
            token_input.setText(current_token)
        layout.addWidget(token_input)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dlg.accept)
        buttons.rejected.connect(dlg.reject)
        layout.addWidget(buttons)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return False
        token = token_input.text().strip()
        if not token.startswith("hf_"):
            QMessageBox.warning(self, self._t("Неверный токен", "Invalid token"), self._t("Токен должен начинаться с 'hf_'", "The token must start with 'hf_'."))
            return False
        try:
            env_path = save_env_value("HF_TOKEN", token)
            self.log(f"Токен сохранён: {env_path}")
        except Exception as exc:
            QMessageBox.warning(
                self,
                self._t("Не удалось сохранить токен", "Could not save token"),
                str(exc),
            )
            return False
        return True

    def _toggle_diarization(self, state):
        if self._diarization_prompt_open:
            return
        enabling = (state == Qt.CheckState.Checked.value)
        if enabling and not os.getenv("HF_TOKEN", "").startswith("hf_"):
            self._diarization_prompt_open = True
            self.cb_diarization.setEnabled(False)
            try:
                if not self._show_hf_token_dialog():
                    self.cb_diarization.blockSignals(True)
                    self.cb_diarization.setChecked(False)
                    self.cb_diarization.blockSignals(False)
                    self.enable_diarization = False
                    self.entry_num_speakers.setEnabled(False)
                    for fmt in ('txt_diarize', 'txt_diarize_timecodes'):
                        cb = self.format_checkboxes.get(fmt)
                        if cb:
                            cb.setEnabled(False)
                    return
            finally:
                self.cb_diarization.setEnabled(True)
                self._diarization_prompt_open = False
        self.enable_diarization = enabling
        self.entry_num_speakers.setEnabled(self.enable_diarization)
        for fmt in ('txt_diarize', 'txt_diarize_timecodes'):
            cb = self.format_checkboxes.get(fmt)
            if cb:
                cb.setEnabled(self.enable_diarization)
        if self.enable_diarization:
            self.log("Диаризация спикеров: ВКЛЮЧЕНА")
        else:
            self.entry_num_speakers.setValue(0)
            self.log("Диаризация спикеров: ВЫКЛЮЧЕНА")

    def _toggle_format(self, fmt: str):
        self.output_formats[fmt] = self.format_checkboxes[fmt].isChecked()

    def _get_selected_formats(self) -> list:
        return [fmt for fmt, enabled in self.output_formats.items() if enabled]

    # ──────────────────────────────────────────────────────────────
    # Файлы / папки
    # ──────────────────────────────────────────────────────────────

    def log(self, message: str):
        message = self._translate_runtime_text(message)
        self.signals.log_message.emit(message)
        self.app_logger.get_logger().info(message)

    def _append_log(self, message: str):
        self.log_text.append(f">> {message}")

    def open_paths_from_system(self, paths: list, append: bool = True):
        """Open files received from Finder, Dock, CLI args, or another app instance."""
        media_files, transcript_files = self._collect_supported_open_paths(paths)
        if media_files:
            if hasattr(self, "tabs"):
                self.tabs.setCurrentIndex(0)
            self._apply_dropped_or_selected_files(media_files, append=append)
        if transcript_files:
            if hasattr(self, "tabs"):
                self.tabs.setCurrentIndex(1)
            self.transcript_files_for_llm = transcript_files if not append else self._merge_paths(
                self.transcript_files_for_llm, transcript_files
            )
            folder = os.path.dirname(transcript_files[0])
            self.llm_transcript_dir = folder
            self.user_settings.set_value("llm_transcript_dir", folder)
            self.user_settings.set_value("last_selected_transcript_files", self.transcript_files_for_llm)
            self._refresh_llm_files_list()
            self.lbl_llm_status.setText(self._t("Транскрипты готовы к LLM-обработке", "Transcripts are ready for LLM processing"))
        if not media_files and not transcript_files and paths:
            QMessageBox.information(
                self,
                self._t("Неподдерживаемый формат", "Unsupported format"),
                self._t("Файлы не являются поддерживаемыми медиа или транскриптами (.txt, .md, .srt, .vtt).", "Files are not supported media or transcript files (.txt, .md, .srt, .vtt)."),
            )
        self.showNormal()
        self.raise_()
        self.activateWindow()

    def _merge_paths(self, existing_paths: list, new_paths: list) -> list:
        merged = []
        seen = set()
        for path in list(existing_paths) + list(new_paths):
            normalized = os.path.abspath(path)
            if normalized not in seen:
                merged.append(path)
                seen.add(normalized)
        return merged

    def _collect_supported_open_paths(self, paths: list):
        transcript_exts = (".txt", ".md", ".srt", ".vtt")
        media_files = []
        transcript_files = []
        for raw_path in paths:
            path = os.path.abspath(os.path.expanduser(str(raw_path)))
            if os.path.isdir(path):
                for root, _dirs, filenames in os.walk(path):
                    for filename in filenames:
                        full = os.path.join(root, filename)
                        lower = filename.lower()
                        if lower.endswith(MEDIA_EXTENSIONS):
                            media_files.append(full)
                        elif lower.endswith(transcript_exts):
                            transcript_files.append(full)
            elif os.path.isfile(path):
                lower = path.lower()
                if lower.endswith(MEDIA_EXTENSIONS):
                    media_files.append(path)
                elif lower.endswith(transcript_exts):
                    transcript_files.append(path)
        return self._merge_paths([], media_files), self._merge_paths([], transcript_files)

    def _select_files(self):
        initial_dir = self.user_settings.get_last_files_dir() or self.input_dir or os.path.expanduser("~")
        files, _ = QFileDialog.getOpenFileNames(
            self, "Выберите аудио или видео файлы", initial_dir,
            "Медиа файлы (*.mp3 *.wav *.m4a *.aac *.flac *.ogg *.mp4 *.avi *.mov *.mkv *.webm *.wma *.qta *.3gp);;Все файлы (*.*)"
        )
        if files:
            self._apply_dropped_or_selected_files(files)

    def _apply_dropped_or_selected_files(self, files: list, append: bool = False, remember_dir: bool = True):
        if not files:
            return
        unique_files = []
        seen = set()
        for file_path in files:
            normalized = os.path.abspath(file_path)
            if normalized not in seen:
                unique_files.append(file_path)
                seen.add(normalized)
        if append:
            existing = {os.path.abspath(f) for f in self.files_to_process}
            for file_path in unique_files:
                normalized = os.path.abspath(file_path)
                if normalized not in existing:
                    self.files_to_process.append(file_path)
                    existing.add(normalized)
        else:
            self.files_to_process = unique_files
        if remember_dir:
            file_dir = os.path.dirname(unique_files[0])
            self.input_dir = file_dir
            self.user_settings.set_last_files_dir(file_dir)
        self._refresh_files_list()
        self.user_settings.set_value("last_selected_audio_files", [p for p in self.files_to_process if os.path.isfile(p)])
        self.log(f"Добавлено в очередь: {len(unique_files)} файлов")
        for f in unique_files:
            self.log(f" + {os.path.basename(f)}")

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            self._set_drop_active(True)
            event.acceptProposedAction()

    def dragLeaveEvent(self, event):
        self._set_drop_active(False)
        super().dragLeaveEvent(event)

    def _set_drop_active(self, active: bool):
        self._drop_active = active
        if hasattr(self, "drop_hint"):
            self._style_drop_hint()

    def dropEvent(self, event: QDropEvent):
        self._set_drop_active(False)
        urls = event.mimeData().urls()
        if not urls:
            event.acceptProposedAction()
            return

        self.open_paths_from_system([url.toLocalFile() for url in urls if url.toLocalFile()], append=True)
        event.acceptProposedAction()

    def _select_files_folder(self):
        initial_dir = self.user_settings.get_last_files_dir() or self.input_dir or os.path.expanduser("~")
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку с аудио/видео файлами", initial_dir)
        if folder:
            self.input_dir = folder
            self.user_settings.set_last_files_dir(folder)
            self._update_input_dir_label(folder)
            files = []
            for root, _dirs, filenames in os.walk(folder):
                for f in filenames:
                    if f.lower().endswith(MEDIA_EXTENSIONS):
                        files.append(os.path.join(root, f))
            if files:
                self.files_to_process = files
                self._refresh_files_list()
                self.user_settings.set_value("last_selected_audio_files", [p for p in self.files_to_process if os.path.isfile(p)])
                self.log(f"Добавлено из папки (включая подпапки): {len(files)} файлов")
                for f in files:
                    self.log(f" + {os.path.relpath(f, folder)}")
            else:
                QMessageBox.information(self, self._t("Информация", "Information"), self._t("В выбранной папке и подпапках нет поддерживаемых файлов", "No supported files were found in the selected folder or its subfolders."))

    def _select_output_folder(self):
        initial_dir = self.user_settings.get_last_output_dir() or self.output_dir or os.path.expanduser("~")
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения результатов", initial_dir)
        if folder:
            self.output_dir = folder
            self.user_settings.set_last_output_dir(folder)
            self._update_output_dir_label(folder)
            self.log(f"Папка для сохранения: {folder}")

    def _update_input_dir_label(self, path: str):
        display_path = path if len(path) < 70 else f"...{path[-70:]}"
        self.lbl_input_folder.setText(self._t(f"Папка источника: {display_path}", f"Source folder: {display_path}"))
        self.lbl_input_folder.setStyleSheet(self._transparent_label_style(self._colors()["text_sub"], font_pt=9))

    def _update_output_dir_label(self, path: str):
        display_path = path if len(path) < 60 else f"...{path[-60:]}"
        self.lbl_output_folder.setText(display_path)
        self.lbl_output_folder.setStyleSheet(self._transparent_label_style(self._colors()["text_sub"]))

    def _update_llm_output_dir_label(self, path: str):
        display_path = path if len(path) < 60 else f"...{path[-60:]}"
        self.lbl_llm_output.setText(display_path)
        self.lbl_llm_output.setStyleSheet(self._transparent_label_style(self._colors()["text_sub"]))

    def _restore_ui_settings(self):
        saved_formats = self.user_settings.get_value("output_formats", {}) or {}
        for fmt, cb in self.format_checkboxes.items():
            cb.blockSignals(True)
            cb.setChecked(bool(saved_formats.get(fmt, self.output_formats.get(fmt, False))))
            cb.blockSignals(False)
            self.output_formats[fmt] = cb.isChecked()

        diarization_enabled = bool(self.user_settings.get_value("enable_diarization", False))
        num_speakers = int(self.user_settings.get_value("num_speakers", 0) or 0)
        self.cb_diarization.blockSignals(True)
        self.cb_diarization.setChecked(diarization_enabled)
        self.cb_diarization.blockSignals(False)
        self.enable_diarization = diarization_enabled
        self.entry_num_speakers.setEnabled(diarization_enabled)
        self.entry_num_speakers.setValue(num_speakers)
        for fmt in ('txt_diarize', 'txt_diarize_timecodes'):
            cb = self.format_checkboxes.get(fmt)
            if cb:
                cb.setEnabled(diarization_enabled)

        provider = self._normalize_llm_provider(self.user_settings.get_value("llm_provider", "API"))
        display_provider = ("Другое" if self._lang == "ru" else "Other") if provider == "Other" else provider
        index = self.combo_llm_provider.findText(display_provider)
        self.combo_llm_provider.setCurrentIndex(index if index >= 0 else 0)
        self.entry_llm_api_url.setText(self.user_settings.get_value("llm_api_url", LLM_API_URL))
        self.entry_llm_api_key.setText(LLM_API_KEY)
        self.entry_llm_model.setText(self.user_settings.get_value("llm_model", LLM_MODEL))
        self.entry_llm_temperature.setText(str(self.user_settings.get_value("llm_temperature", LLM_TEMPERATURE)))
        self.entry_llm_claude_path.setText(self.user_settings.get_value("llm_claude_path", shutil.which("claude") or "claude"))
        self.entry_llm_claude_args.setText(self.user_settings.get_value("llm_claude_args", ""))
        self.entry_llm_codex_path.setText(self.user_settings.get_value("llm_codex_path", shutil.which("codex") or "codex"))
        self.entry_llm_codex_args.setText(self.user_settings.get_value("llm_codex_args", ""))
        self.entry_llm_opencode_path.setText(self.user_settings.get_value("llm_opencode_path", shutil.which("opencode") or "opencode"))
        self.entry_llm_opencode_args.setText(self.user_settings.get_value("llm_opencode_args", ""))
        self.entry_llm_pi_path.setText(self.user_settings.get_value("llm_pi_path", shutil.which("pi") or "pi"))
        self.entry_llm_pi_provider.setText(self.user_settings.get_value("llm_pi_provider", ""))
        self.entry_llm_pi_args.setText(self.user_settings.get_value("llm_pi_args", ""))
        self.entry_llm_other_path.setText(self.user_settings.get_value("llm_other_path", ""))
        self.entry_llm_other_args.setText(self.user_settings.get_value("llm_other_args", ""))
        self._update_llm_provider_fields(self.combo_llm_provider.currentText())
        self.txt_llm_summary_prompt.setPlainText(self.user_settings.get_value("llm_summary_prompt", SUMMARY_PROMPT))
        self.txt_llm_tasks_prompt.setPlainText(self.user_settings.get_value("llm_tasks_prompt", TASKS_PROMPT))
        self.txt_llm_custom_prompt.setPlainText(self.user_settings.get_value("llm_custom_prompt", ""))
        self.txt_llm_transcript.setPlainText(self.user_settings.get_value("llm_manual_transcript", ""))

        tab_index = int(self.user_settings.get_value("active_tab_index", 0) or 0)
        if 0 <= tab_index < self.tabs.count():
            self.tabs.setCurrentIndex(tab_index)

        saved_audio_files = self.user_settings.get_value("last_selected_audio_files", []) or []
        self.files_to_process = [path for path in saved_audio_files if os.path.isfile(path)]
        if self.files_to_process:
            self._refresh_files_list()

        saved_llm_actions = self.user_settings.get_value("llm_actions", {}) or {}
        for key, cb in self.llm_action_checkboxes.items():
            cb.setChecked(bool(saved_llm_actions.get(key, cb.isChecked())))

        saved_llm_exports = self.user_settings.get_value("llm_export_formats", {}) or {}
        for key, cb in self.llm_export_checkboxes.items():
            cb.setChecked(bool(saved_llm_exports.get(key, cb.isChecked())))

        saved_llm_files = self.user_settings.get_value("last_selected_transcript_files", []) or []
        self.transcript_files_for_llm = [path for path in saved_llm_files if os.path.isfile(path)]
        self._refresh_llm_files_list()

    def _save_ui_settings(self):
        self.user_settings.set_value("output_formats", self.output_formats)
        self.user_settings.set_value("enable_diarization", self.cb_diarization.isChecked())
        self.user_settings.set_value("num_speakers", self.entry_num_speakers.value())
        self.user_settings.set_value("llm_provider", self._normalize_llm_provider(self.combo_llm_provider.currentText()))
        self.user_settings.set_value("llm_api_url", self.entry_llm_api_url.text().strip())
        llm_api_key = self.entry_llm_api_key.text().strip()
        if llm_api_key:
            try:
                save_env_value("LLM_API_KEY", llm_api_key)
            except OSError as exc:
                self.log(f"Не удалось сохранить LLM API key: {exc}")
        self.user_settings.set_value("llm_model", self.entry_llm_model.text().strip())
        self.user_settings.set_value("llm_temperature", self.entry_llm_temperature.text().strip())
        self.user_settings.set_value("llm_claude_path", self.entry_llm_claude_path.text().strip())
        self.user_settings.set_value("llm_claude_args", self.entry_llm_claude_args.text().strip())
        self.user_settings.set_value("llm_codex_path", self.entry_llm_codex_path.text().strip())
        self.user_settings.set_value("llm_codex_args", self.entry_llm_codex_args.text().strip())
        self.user_settings.set_value("llm_opencode_path", self.entry_llm_opencode_path.text().strip())
        self.user_settings.set_value("llm_opencode_args", self.entry_llm_opencode_args.text().strip())
        self.user_settings.set_value("llm_pi_path", self.entry_llm_pi_path.text().strip())
        self.user_settings.set_value("llm_pi_provider", self.entry_llm_pi_provider.text().strip())
        self.user_settings.set_value("llm_pi_args", self.entry_llm_pi_args.text().strip())
        self.user_settings.set_value("llm_other_path", self.entry_llm_other_path.text().strip())
        self.user_settings.set_value("llm_other_args", self.entry_llm_other_args.text().strip())
        self.user_settings.set_value("llm_summary_prompt", self.txt_llm_summary_prompt.toPlainText())
        self.user_settings.set_value("llm_tasks_prompt", self.txt_llm_tasks_prompt.toPlainText())
        self.user_settings.set_value("llm_custom_prompt", self.txt_llm_custom_prompt.toPlainText())
        self.user_settings.set_value("llm_manual_transcript", self.txt_llm_transcript.toPlainText())
        self.user_settings.set_value("active_tab_index", self.tabs.currentIndex())
        self.user_settings.set_value("llm_actions", {k: cb.isChecked() for k, cb in self.llm_action_checkboxes.items()})
        self.user_settings.set_value("llm_export_formats", {k: cb.isChecked() for k, cb in self.llm_export_checkboxes.items()})
        self.user_settings.set_value("last_selected_audio_files", [p for p in self.files_to_process if os.path.isfile(p)])
        self.user_settings.set_value("last_selected_transcript_files", [p for p in self.transcript_files_for_llm if os.path.isfile(p)])
        if self.llm_output_dir:
            self.user_settings.set_value("llm_output_dir", self.llm_output_dir)
        if self.llm_transcript_dir:
            self.user_settings.set_value("llm_transcript_dir", self.llm_transcript_dir)

    def _refresh_llm_files_list(self):
        if not hasattr(self, "llm_files_list"):
            return
        self.llm_files_list.clear()
        for path in self.transcript_files_for_llm:
            item = QListWidgetItem(os.path.basename(path))
            item.setToolTip(path)
            item.setData(Qt.ItemDataRole.UserRole, path)
            self.llm_files_list.addItem(item)
        has_files = bool(self.transcript_files_for_llm)
        self.llm_files_list.setVisible(has_files)
        self.llm_drop_hint.setVisible(not has_files)
        if has_files:
            self._size_list_to_contents(self.llm_files_list)
        c = self._colors()
        if has_files:
            text = self._t(f"Выбрано транскриптов: {len(self.transcript_files_for_llm)}", f"Selected transcripts: {len(self.transcript_files_for_llm)}")
            self.lbl_llm_files.setText(text)
            self.lbl_llm_files_count.setText(text)
            self.lbl_llm_files.setStyleSheet(self._transparent_label_style(c["text_sub"]))
            self.lbl_llm_files_count.setStyleSheet(self._transparent_label_style(c["text_sub"]))
        else:
            text = self._t("Файлы не выбраны", "No files selected")
            self.lbl_llm_files.setText(text)
            self.lbl_llm_files_count.setText(text)
            self.lbl_llm_files.setStyleSheet(self._transparent_label_style(c["text_mute"]))
            self.lbl_llm_files_count.setStyleSheet(self._transparent_label_style(c["text_mute"]))
        self._update_llm_files_controls()

    def _update_llm_files_controls(self):
        if not hasattr(self, "btn_clear_llm_files"):
            return
        has_files = bool(self.transcript_files_for_llm)
        self.btn_clear_llm_files.setEnabled(has_files and not self.is_llm_processing)
        self.btn_remove_llm_file.setEnabled(bool(self.llm_files_list.selectedItems()) and not self.is_llm_processing)

    def _remove_selected_llm_files(self):
        if self.is_llm_processing:
            return
        selected = {item.data(Qt.ItemDataRole.UserRole) for item in self.llm_files_list.selectedItems()}
        if not selected:
            return
        self.transcript_files_for_llm = [p for p in self.transcript_files_for_llm if p not in selected]
        self._refresh_llm_files_list()
        self.user_settings.set_value("last_selected_transcript_files", [p for p in self.transcript_files_for_llm if os.path.isfile(p)])

    def _clear_llm_files_list(self):
        if self.is_llm_processing or not self.transcript_files_for_llm:
            return
        self.transcript_files_for_llm = []
        self._refresh_llm_files_list()
        self.user_settings.set_value("last_selected_transcript_files", [])

    def _select_llm_transcript_files(self):
        initial_dir = self.user_settings.get_value("llm_transcript_dir", self.llm_transcript_dir)
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Выберите транскрипты",
            initial_dir,
            "Транскрипты (*.txt *.md *.srt *.vtt);;Текстовые файлы (*.txt *.md);;Все файлы (*.*)"
        )
        if files:
            self.transcript_files_for_llm = files
            folder = os.path.dirname(files[0])
            self.llm_transcript_dir = folder
            if not self.llm_output_dir:
                self.llm_output_dir = folder
                self._update_llm_output_dir_label(folder)
            self.user_settings.set_value("llm_transcript_dir", folder)
            self.user_settings.set_value("last_selected_transcript_files", files)
            self._refresh_llm_files_list()
            self.lbl_llm_status.setText(self._t("Транскрипты готовы к LLM-обработке", "Transcripts are ready for LLM processing"))

    def _select_llm_output_folder(self):
        initial_dir = self.llm_output_dir or self.output_dir or os.path.expanduser("~")
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения LLM-результатов", initial_dir)
        if folder:
            self.llm_output_dir = folder
            self.user_settings.set_value("llm_output_dir", folder)
            self._update_llm_output_dir_label(folder)

    # ──────────────────────────────────────────────────────────────
    # Загрузка по ссылке
    # ──────────────────────────────────────────────────────────────

    def _start_download(self, start_after_download: bool = False):
        url = self.input_path.text().strip()
        if not url:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), self._t("Введите ссылку для загрузки.", "Enter a URL to download."))
            return
        if not (url.startswith("http://") or url.startswith("https://")):
            QMessageBox.warning(self, self._t("Внимание", "Attention"), self._t("Ссылка должна начинаться с http:// или https://.", "The URL must start with http:// or https://."))
            return
        if self.is_processing:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), self._t("Дождитесь завершения обработки файлов.", "Wait for file processing to finish."))
            return
        if self.is_downloading:
            QMessageBox.information(self, self._t("Информация", "Information"), self._t("Загрузка уже выполняется.", "A download is already in progress."))
            return
        download_dir = self.input_dir
        if not download_dir:
            initial_dir = self.user_settings.get_last_files_dir() or os.path.expanduser("~")
            download_dir = QFileDialog.getExistingDirectory(self, "Выберите папку для загрузки медиа", initial_dir)
            if not download_dir:
                return
            self.input_dir = download_dir
            self.user_settings.set_last_files_dir(download_dir)
            self._update_input_dir_label(download_dir)
        self.is_downloading = True
        self.start_processing_after_download = start_after_download
        self.btn_upload.setEnabled(False)
        self.btn_start.setEnabled(False)
        self.progress_upload.setValue(0)
        self.progress_upload.setVisible(True)
        self.lbl_status.setText("Загрузка медиа по ссылке...")
        self._set_status("Загрузка по ссылке…")
        self.log(f"Загрузка медиа по ссылке в папку: {download_dir}")
        threading.Thread(target=self._download_media, args=(url, download_dir), daemon=True).start()

    def _download_media(self, url: str, download_dir: str):
        try:
            result = self.media_downloader.download(
                url, download_dir,
                progress_callback=self.signals.download_progress.emit,
                allow_playlist=False,
            )
            files = [p for p in result.files if os.path.isfile(p) and os.path.getsize(p) > 0]
            if not files:
                raise RuntimeError("yt-dlp не вернул скачанный медиафайл")
            self.signals.download_finished.emit(files)
        except Exception as e:
            self.signals.download_failed.emit(str(e))

    def _update_download_progress(self, value: int):
        self.progress_upload.setValue(value)

    def _on_download_finished(self, files: list):
        self.is_downloading = False
        self.btn_upload.setEnabled(True)
        self.btn_start.setEnabled(True)
        self.progress_upload.setValue(0)
        self.progress_upload.setVisible(False)
        if files:
            self._apply_dropped_or_selected_files(files, append=True, remember_dir=False)
            self.input_path.clear()
            self.lbl_status.setText(self._t("Медиа загружено и добавлено в очередь", "Media downloaded and added to the queue"))
            self.log(f"Загрузка завершена: {len(files)} файлов")
        else:
            QMessageBox.warning(self, self._t("Загрузка", "Download"), self._t("Не удалось получить медиафайлы по ссылке.", "Failed to get media files from the URL."))
            self.log("Загрузка завершилась без файлов")
        if self.start_processing_after_download:
            self.start_processing_after_download = False
            QTimer.singleShot(0, self._start_processing_thread)

    def _on_download_failed(self, message: str):
        self.is_downloading = False
        self.start_processing_after_download = False
        self.btn_upload.setEnabled(True)
        self.btn_start.setEnabled(True)
        self.progress_upload.setValue(0)
        self.progress_upload.setVisible(False)
        self.lbl_status.setText(self._t("Ошибка загрузки", "Download error"))
        self.log(f"Ошибка загрузки: {message}")
        QMessageBox.warning(self, self._t("Ошибка загрузки", "Download error"), message)

    def _clear_llm_result(self):
        if self.is_llm_processing:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("LLM-обработка уже выполняется", "LLM processing is already running."))
            return
        self.txt_llm_result.clear()
        self.llm_last_result_text = ""
        self.llm_last_result_name = "llm_result"
        self.lbl_llm_status.setText(self._t("Готово к LLM-обработке", "Ready for LLM processing"))

    def _clear_llm_all(self):
        if self.is_llm_processing:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("LLM-обработка уже выполняется", "LLM processing is already running."))
            return
        self.transcript_files_for_llm = []
        self.txt_llm_transcript.clear()
        self._refresh_llm_files_list()
        self._clear_llm_result()
        self.user_settings.set_value("last_selected_transcript_files", [])
        self.user_settings.set_value("llm_manual_transcript", "")

    def _selected_llm_modes(self):
        modes = []
        if self.llm_action_checkboxes["summary"].isChecked():
            prompt = self.txt_llm_summary_prompt.toPlainText().strip() or SUMMARY_PROMPT
            modes.append(("summary", "Выжимка", prompt))
        if self.llm_action_checkboxes["tasks"].isChecked():
            prompt = self.txt_llm_tasks_prompt.toPlainText().strip() or TASKS_PROMPT
            modes.append(("tasks", "Задачи", prompt))
        if self.llm_action_checkboxes["custom"].isChecked():
            custom_prompt = self.txt_llm_custom_prompt.toPlainText().strip()
            if not custom_prompt:
                raise ValueError("Для режима «Свой промпт» укажите пользовательский промпт в меню «Настройки → LLM API…»")
            modes.append(("custom", "Свой промпт", custom_prompt))
        if not modes:
            raise ValueError("Выберите хотя бы один чекбокс в блоке «Что делать»")
        return modes

    def _selected_llm_export_formats(self):
        formats = [key for key, cb in self.llm_export_checkboxes.items() if cb.isChecked()]
        if not formats:
            raise ValueError("Выберите хотя бы один формат сохранения результата")
        return formats

    def _start_llm_processing(self):
        if self.is_llm_processing:
            return
        try:
            llm_settings = self._collect_llm_settings()
            items = self._collect_llm_inputs()
            modes = self._selected_llm_modes()
            export_formats = self._selected_llm_export_formats()
        except ValueError as e:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), str(e))
            return

        self._save_ui_settings()
        self.is_llm_processing = True
        self._set_llm_buttons_enabled(False)
        self.lbl_llm_status.setText("Идет LLM-обработка...")
        self.txt_llm_result.clear()
        threading.Thread(
            target=self._run_llm_processing,
            args=(llm_settings, items, modes, export_formats),
            daemon=True,
        ).start()

    def _collect_llm_settings(self) -> dict:
        provider = self.combo_llm_provider.currentText().strip() or "API"
        api_url = self.entry_llm_api_url.text().strip()
        api_key = self.entry_llm_api_key.text().strip()
        model = self.entry_llm_model.text().strip()
        temperature_text = self.entry_llm_temperature.text().strip() or str(LLM_TEMPERATURE)
        try:
            temperature = float(temperature_text)
        except ValueError:
            raise ValueError("Temperature должно быть числом")
        if not 0 <= temperature <= 2:
            raise ValueError("Temperature должно быть в диапазоне 0..2")

        if provider == "API":
            if not api_url:
                raise ValueError("Укажите API URL")
            if not api_key:
                raise ValueError("Укажите API Key")
            if not model:
                raise ValueError("Укажите модель")
        elif provider == "Claude Code":
            claude_path = self.entry_llm_claude_path.text().strip() or "claude"
            if not (shutil.which(claude_path) or os.path.isfile(claude_path)):
                raise ValueError(f"Не найден Claude Code: {claude_path}")
        elif provider == "Codex":
            codex_path = self.entry_llm_codex_path.text().strip() or "codex"
            if not (shutil.which(codex_path) or os.path.isfile(codex_path)):
                raise ValueError(f"Не найден Codex: {codex_path}")
        elif provider == "OpenCode":
            opencode_path = self.entry_llm_opencode_path.text().strip() or "opencode"
            if not (shutil.which(opencode_path) or os.path.isfile(opencode_path)):
                raise ValueError(f"Не найден OpenCode: {opencode_path}")
        elif provider == "Pi":
            pi_path = self.entry_llm_pi_path.text().strip() or "pi"
            if not (shutil.which(pi_path) or os.path.isfile(pi_path)):
                raise ValueError(f"Не найден Pi: {pi_path}")
        elif self._normalize_llm_provider(provider) == "Other":
            other_path = self.entry_llm_other_path.text().strip()
            if not other_path:
                raise ValueError(self._t("Укажите команду для провайдера «Другое»", "Specify a command for the 'Other' provider"))
            if not (shutil.which(other_path) or os.path.isfile(other_path)):
                raise ValueError(f"Не найдена команда: {other_path}")
        return {
            "provider": provider,
            "api_url": api_url,
            "api_key": api_key,
            "model": model,
            "temperature": temperature,
            "claude_path": self.entry_llm_claude_path.text().strip() or "claude",
            "claude_args": self.entry_llm_claude_args.text().strip(),
            "codex_path": self.entry_llm_codex_path.text().strip() or "codex",
            "codex_args": self.entry_llm_codex_args.text().strip(),
            "opencode_path": self.entry_llm_opencode_path.text().strip() or "opencode",
            "opencode_args": self.entry_llm_opencode_args.text().strip(),
            "pi_path": self.entry_llm_pi_path.text().strip() or "pi",
            "pi_provider": self.entry_llm_pi_provider.text().strip(),
            "pi_args": self.entry_llm_pi_args.text().strip(),
            "other_path": self.entry_llm_other_path.text().strip(),
            "other_args": self.entry_llm_other_args.text().strip(),
        }

    def _collect_llm_inputs(self):
        manual_text = self.txt_llm_transcript.toPlainText().strip()
        items = []
        if manual_text:
            base_name = "manual_transcript"
            if self.transcript_files_for_llm:
                base_name = Path(self.transcript_files_for_llm[0]).stem
            items.append({"name": base_name, "text": manual_text, "source_path": self.transcript_files_for_llm[0] if self.transcript_files_for_llm else None})
        for path in self.transcript_files_for_llm:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read().strip()
            except OSError:
                continue
            if text:
                items.append({"name": Path(path).stem, "text": text, "source_path": path})
        if not items:
            raise ValueError("Выберите хотя бы один транскрипт или вставьте текст вручную")
        return items

    def _build_llm_prompt_text(self, transcript_text: str, prompt: str) -> str:
        return (
            "Ты обрабатываешь транскрипт на русском языке. "
            "Не выдумывай факты, явно помечай неясности.\n\n"
            f"Инструкция:\n{prompt.strip()}\n\n"
            f"Транскрипт:\n{transcript_text.strip()}\n"
        )

    def _run_api_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        client = LLMClient(LLMSettings(
            api_url=llm_settings["api_url"],
            api_key=llm_settings["api_key"],
            model=llm_settings["model"],
            temperature=llm_settings["temperature"],
        ))
        return client.process_transcript(transcript_text, prompt)

    def _run_claude_code_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        command = [llm_settings["claude_path"], "-p", "--output-format", "text"]
        if llm_settings.get("model"):
            command += ["--model", llm_settings["model"]]
        if llm_settings.get("claude_args"):
            command += shlex.split(llm_settings["claude_args"])
        command.append(self._build_llm_prompt_text(transcript_text, prompt))
        result = subprocess.run(command, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            raise RuntimeError((result.stderr or result.stdout or "Claude Code завершился с ошибкой").strip())
        answer = (result.stdout or "").strip()
        if not answer:
            raise RuntimeError("Claude Code вернул пустой ответ")
        return answer

    def _run_codex_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tmp:
            output_path = tmp.name
        try:
            command = [llm_settings["codex_path"], "exec", "-o", output_path]
            if llm_settings.get("model"):
                command += ["-m", llm_settings["model"]]
            if llm_settings.get("codex_args"):
                command += shlex.split(llm_settings["codex_args"])
            command.append("-")
            result = subprocess.run(
                command,
                input=self._build_llm_prompt_text(transcript_text, prompt),
                capture_output=True,
                text=True,
                timeout=600,
            )
            if result.returncode != 0:
                raise RuntimeError((result.stderr or result.stdout or "Codex завершился с ошибкой").strip())
            with open(output_path, "r", encoding="utf-8") as f:
                answer = f.read().strip()
            if not answer:
                raise RuntimeError("Codex вернул пустой ответ")
            return answer
        finally:
            try:
                os.remove(output_path)
            except OSError:
                pass

    def _run_generic_cli_prompt(self, command: list[str], error_name: str) -> str:
        result = subprocess.run(command, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            raise RuntimeError((result.stderr or result.stdout or f"{error_name} завершился с ошибкой").strip())
        answer = (result.stdout or "").strip()
        if not answer:
            raise RuntimeError(f"{error_name} вернул пустой ответ")
        return answer

    def _run_opencode_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        command = [llm_settings["opencode_path"]]
        if llm_settings.get("model"):
            command += ["--model", llm_settings["model"]]
        if llm_settings.get("opencode_args"):
            command += shlex.split(llm_settings["opencode_args"])
        command.append(self._build_llm_prompt_text(transcript_text, prompt))
        return self._run_generic_cli_prompt(command, "OpenCode")

    def _run_pi_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        command = [llm_settings["pi_path"], "-p", "--mode", "text"]
        if llm_settings.get("pi_provider"):
            command += ["--provider", llm_settings["pi_provider"]]
        if llm_settings.get("model"):
            command += ["--model", llm_settings["model"]]
        if llm_settings.get("pi_args"):
            command += shlex.split(llm_settings["pi_args"])
        command.append(self._build_llm_prompt_text(transcript_text, prompt))
        return self._run_generic_cli_prompt(command, "Pi")

    def _run_other_llm(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        command = [llm_settings["other_path"]]
        if llm_settings.get("other_args"):
            command += shlex.split(llm_settings["other_args"])
        command.append(self._build_llm_prompt_text(transcript_text, prompt))
        return self._run_generic_cli_prompt(command, "Внешний CLI")

    def _run_llm_provider(self, llm_settings: dict, transcript_text: str, prompt: str) -> str:
        provider = llm_settings.get("provider", "API")
        if provider == "API":
            return self._run_api_llm(llm_settings, transcript_text, prompt)
        if provider == "Claude Code":
            return self._run_claude_code_llm(llm_settings, transcript_text, prompt)
        if provider == "Codex":
            return self._run_codex_llm(llm_settings, transcript_text, prompt)
        if provider == "OpenCode":
            return self._run_opencode_llm(llm_settings, transcript_text, prompt)
        if provider == "Pi":
            return self._run_pi_llm(llm_settings, transcript_text, prompt)
        if self._normalize_llm_provider(provider) == "Other":
            return self._run_other_llm(llm_settings, transcript_text, prompt)
        raise RuntimeError(self._t(f"Неизвестный LLM-провайдер: {provider}", f"Unknown LLM provider: {provider}"))

    def _run_llm_processing(self, llm_settings: dict, items: list, modes: list, export_formats: list):
        try:
            results = []
            total = len(items)
            last_name = "llm_result"
            provider = llm_settings.get("provider", "API")
            for item_index, item in enumerate(items, start=1):
                name = item["name"]
                last_name = name
                item_blocks = []
                for mode_suffix, mode_label, prompt in modes:
                    self.log(f"LLM: обработка {item_index}/{total} — {name} — {mode_label} — {provider}")
                    answer = self._run_llm_provider(llm_settings, item["text"], prompt)
                    saved_paths = self._save_llm_result(item, answer, mode_suffix, export_formats)
                    block = f"=== {name} / {mode_label} / {provider} ===\n{answer}"
                    if saved_paths:
                        block += "\n\nСохранено:\n" + "\n".join(saved_paths)
                    item_blocks.append(block)
                results.append("\n\n".join(item_blocks))
            mode_suffixes = "_".join(mode[0] for mode in modes)
            self.llm_last_result_name = f"{last_name}_llm_{mode_suffixes}"
            final_text = "\n\n".join(results)
            self.signals.llm_finished.emit(True, f"LLM-обработка завершена: {total} файл(ов)", final_text)
        except Exception as e:
            error_text = str(e).strip() or "Неизвестная ошибка"
            self.signals.llm_finished.emit(False, f"Ошибка LLM: {self._compact_llm_error(error_text)}", error_text)

    def _compact_llm_error(self, error_text: str, limit: int = 180) -> str:
        raw_text = (error_text or "").strip()
        lowered = raw_text.lower()

        friendly_rules = [
            (("refresh token was revoked", "please log out and sign in again"), "Codex: сессия истекла или токен отозван — нужно заново войти в Codex"),
            (("token_invalidated", "authentication token has been invalidated"), "Codex: токен недействителен — перелогиньтесь"),
            (("your session has ended", "refresh_token_invalidated"), "Codex: сессия завершилась — выполните codex logout и codex login"),
            (("connection refused", "127.0.0.1", "/v1/responses"), "Codex: локальный backend недоступен — проверьте, запущен ли нужный сервер/провайдер"),
            (("failed to refresh available models", "missing field `base_instructions`"), "Codex: сервер моделей отдает несовместимый формат ответа — провайдер/прокси не полностью совместим с Codex"),
            (("failed to decode models response",), "Codex: провайдер вернул неожиданный формат списка моделей"),
            (("401", "anthropic"), "Anthropic API: ошибка авторизации (401) — проверьте API key"),
            (("401", "openai"), "OpenAI-compatible API: ошибка авторизации (401) — проверьте API key"),
            (("401", "unauthorized"), "Ошибка авторизации (401) — проверьте ключ, токен или логин выбранного провайдера"),
            (("403", "forbidden"), "Доступ запрещен (403) — у аккаунта или ключа не хватает прав"),
            (("404",), "Endpoint не найден (404) — проверьте URL API и путь /v1/..."),
            (("429", "rate"), "Превышен лимит запросов (429) — попробуйте позже или смените тариф/провайдера"),
            (("insufficient_quota",), "Закончилась квота API — проверьте биллинг или лимиты"),
            (("model_not_found",), "Указанная модель не найдена — проверьте точное имя модели"),
            (("does not exist", "model"), "Указанная модель не существует у выбранного провайдера"),
            (("invalid x-api-key",), "Неверный Anthropic API key"),
            (("incorrect api key",), "Неверный API key"),
            (("could not resolve host",), "Не удалось найти хост — проверьте URL и интернет-соединение"),
            (("name or service not known",), "Не удалось найти сервер — проверьте адрес API"),
            (("max retries exceeded",), "Не удалось подключиться к API после нескольких попыток"),
            (("read timed out", "timeout"), "Сервер слишком долго отвечает — попробуйте позже или увеличьте timeout"),
            (("connection timed out",), "Таймаут соединения — сервер недоступен или отвечает слишком долго"),
            (("ssl", "certificate"), "Ошибка SSL-сертификата — проверьте HTTPS/сертификат сервера"),
            (("command not found",), "Не найдена команда CLI-провайдера — проверьте путь в настройках"),
            (("not found", "claude"), "Claude Code не найден — проверьте путь к команде claude"),
            (("not found", "codex"), "Codex не найден — проверьте путь к команде codex"),
            (("not found", "opencode"), "OpenCode не найден — проверьте путь к команде opencode"),
            (("not found", "pi"), "Pi не найден — проверьте путь к команде pi"),
        ]
        for needles, message in friendly_rules:
            if all(needle in lowered for needle in needles):
                return message

        text = " ".join(raw_text.split())
        if len(text) <= limit:
            return text
        return text[:limit - 1] + "…"

    def _write_llm_output_file(self, save_path: str, content: str, export_format: str):
        if export_format in ("txt", "md"):
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(content)
            return
        if export_format == "docx":
            from docx import Document
            doc = Document()
            for block in content.split("\n\n"):
                doc.add_paragraph(block)
            doc.save(save_path)

    def _save_llm_result(self, item: dict, answer: str, mode_suffix: str, export_formats: list) -> list[str]:
        source_path = item.get("source_path")
        if self.llm_output_dir:
            target_dir = self.llm_output_dir
        elif source_path:
            target_dir = os.path.dirname(source_path)
        else:
            target_dir = os.getcwd()
        os.makedirs(target_dir, exist_ok=True)
        saved_paths = []
        for export_format in export_formats:
            save_path = os.path.join(target_dir, f"{item['name']}_llm_{mode_suffix}.{export_format}")
            self._write_llm_output_file(save_path, answer, export_format)
            saved_paths.append(save_path)
        return saved_paths

    def _set_llm_buttons_enabled(self, enabled: bool):
        self.btn_llm_process.setEnabled(enabled)
        self.btn_llm_clear.setEnabled(enabled)

    def _on_llm_finished(self, success: bool, message: str, result_text: str):
        self.is_llm_processing = False
        self._set_llm_buttons_enabled(True)
        self.lbl_llm_status.setText(message)
        if result_text:
            self.llm_last_result_text = result_text
            self.txt_llm_result.setPlainText(result_text)
        elif not success:
            self.llm_last_result_text = ""
            self.txt_llm_result.setPlainText(message)
        if success:
            QMessageBox.information(self, self._t("Готово", "Done"), message)
        else:
            QMessageBox.warning(self, self._t("Ошибка", "Error"), message)

    def _export_llm_result(self, export_format: str):
        result_text = self.txt_llm_result.toPlainText().strip()
        if not result_text:
            QMessageBox.information(self, self._t("Внимание", "Attention"), self._t("Нет результата для экспорта", "There is no result to export."))
            return
        suffix = {"txt": "txt", "md": "md", "docx": "docx"}[export_format]
        initial_dir = self.llm_output_dir or self.output_dir or os.path.expanduser("~")
        default_name = f"{self.llm_last_result_name}.{suffix}"
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            f"Сохранить результат как {suffix.upper()}",
            os.path.join(initial_dir, default_name),
            f"{suffix.upper()} files (*.{suffix})"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(f".{suffix}"):
            save_path += f".{suffix}"
        try:
            if export_format in ("txt", "md"):
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(result_text)
            else:
                try:
                    from docx import Document
                except Exception:
                    QMessageBox.warning(self, self._t("Ошибка", "Error"), self._t("Для экспорта в DOCX установите пакет python-docx", "Install the python-docx package to export to DOCX."))
                    return
                doc = Document()
                for block in result_text.split("\n\n"):
                    doc.add_paragraph(block)
                doc.save(save_path)
            target_dir = os.path.dirname(save_path)
            if target_dir:
                self.llm_output_dir = target_dir
                self.user_settings.set_value("llm_output_dir", target_dir)
                self._update_llm_output_dir_label(target_dir)
            self.lbl_llm_status.setText(f"Результат экспортирован: {os.path.basename(save_path)}")
        except Exception as e:
            QMessageBox.warning(self, self._t("Ошибка", "Error"), self._t("Не удалось экспортировать результат: ", "Failed to export the result: ") + str(e))

    # ──────────────────────────────────────────────────────────────
    # Обработка файлов
    # ──────────────────────────────────────────────────────────────

    def _clear_all(self):
        if self.is_processing:
            reply = QMessageBox.question(
                self,
                self._t("Внимание", "Attention"),
                self._t("Идет обработка файлов. Вы уверены, что хотите сбросить все настройки?", "File processing is in progress. Are you sure you want to reset all settings?"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                return
            self._cancel_requested = True
            self.is_processing = False
            self.progress_timer.stop()
            self._set_processing_controls_enabled(True)
        self.files_to_process = []
        self.output_dir = ""
        self.input_dir = ""
        self.files_processed = 0
        self.total_files = 0
        self.time_spent = 0
        self.current_file_start_time = 0
        self.start_time = None
        self.file_estimates = {}
        self.total_estimated_time = 0
        self.current_stage = None
        self.current_stage_progress = 0.0
        self.start_processing_after_download = False
        self._last_result_dir = ""
        self.log_text.clear()
        self.progress_bar_total.setValue(0)
        self.progress_bar_file.setValue(0)
        self.progress_upload.setValue(0)
        self.progress_upload.setVisible(False)
        self.lbl_current_file.setText("")
        self.lbl_stage.setText("")
        self.lbl_file_counter.setText("")
        self.detail_row.setVisible(False)
        self.lbl_status.setText(self._t("Готов к работе", "Ready to work"))
        self.btn_cancel.setVisible(False)
        self.btn_cancel.setEnabled(True)
        self.btn_cancel.setText(self._t("Отменить", "Cancel"))
        self.btn_open_result.setVisible(False)
        c = self._colors()
        self._refresh_files_list()
        self.lbl_input_folder.setText(self._t("Папка не выбрана", "Folder not selected"))
        self.lbl_input_folder.setStyleSheet(self._transparent_label_style(c["text_mute"], font_pt=9))
        self.lbl_output_folder.setText(self._t("Папка не выбрана (по умолчанию - рядом с файлом)", "Folder not selected (default: next to the file)"))
        self.lbl_output_folder.setStyleSheet(self._transparent_label_style(c["text_mute"]))
        self.input_path.clear()
        self.btn_start.setEnabled(True)
        self.btn_upload.setEnabled(True)
        self._set_status(self._t("Готов к работе", "Ready to work"))
        self.log(self._t("Все настройки сброшены", "All settings have been reset"))

    def _cancel_processing(self):
        if not self.is_processing or self._cancel_requested:
            return
        self._cancel_requested = True
        self.btn_cancel.setEnabled(False)
        self.btn_cancel.setText(self._t("Отмена…", "Cancelling…"))
        self.lbl_stage.setText(self._t("●  Останавливаем после текущего файла…", "●  Stopping after the current file…"))
        self.lbl_status.setText(self._t("Отмена: дождитесь завершения текущего файла…", "Cancellation: wait for the current file to finish…"))
        self._set_status(self._t("Отмена обработки…", "Cancelling processing…"))
        self.log(self._t("Запрошена отмена обработки — остановимся после текущего файла", "Cancellation requested — stopping after the current file"))

    def _start_processing_thread(self):
        if self.is_processing:
            return
        if self.is_downloading:
            QMessageBox.information(self, self._t("Информация", "Information"), self._t("Дождитесь завершения загрузки по ссылке.", "Wait for the URL download to finish."))
            return
        if self.input_path.text().strip():
            self._start_download(start_after_download=True)
            return
        if not self.files_to_process:
            QMessageBox.warning(self, self._t("Внимание", "Attention"), self._t("Выберите хотя бы один файл для обработки!", "Choose at least one file for processing!"))
            return
        if not any(self.output_formats.values()):
            QMessageBox.warning(self, self._t("Внимание", "Attention"), self._t("Выберите хотя бы один формат вывода!", "Choose at least one output format!"))
            return
        if not self.output_dir:
            self.log(self._t("Папка сохранения не выбрана. Результаты будут сохраняться рядом с каждым исходным файлом.", "Output folder not selected. Results will be saved next to each source file."))
        self.is_processing = True
        self._cancel_requested = False
        self.start_time = time.time()
        self.files_processed = 0
        self.total_files = len(self.files_to_process)
        self.time_spent = 0
        self.current_file_start_time = 0
        self.current_stage = None
        self.current_stage_progress = 0.0
        self.log(self._t("Анализ файлов и оценка времени обработки...", "Analyzing files and estimating processing time..."))
        files_with_durations = []
        for filepath in self.files_to_process:
            try:
                duration = AudioConverter.get_media_duration(filepath)
                if duration > 0:
                    self.log(f"  {os.path.basename(filepath)}: {int(duration//60)}:{int(duration%60):02d}")
                    files_with_durations.append((filepath, duration))
                else:
                    self.log(self._t(f"  {os.path.basename(filepath)}: длительность неизвестна", f"  {os.path.basename(filepath)}: duration unknown"))
                    files_with_durations.append((filepath, 60))
            except Exception:
                self.log(self._t(f"  {os.path.basename(filepath)}: ошибка определения длительности", f"  {os.path.basename(filepath)}: duration detection error"))
                files_with_durations.append((filepath, 60))
        batch_estimate = self.stats.estimate_batch_time(files_with_durations)
        self.file_estimates = batch_estimate["per_file"]
        self.total_estimated_time = batch_estimate["total_seconds"]
        estimate_str = self.time_formatter.format_duration(self.total_estimated_time)
        self.log(self._t(f"Ожидаемое время обработки: ~{estimate_str}", f"Estimated processing time: ~{estimate_str}"))
        self.btn_start.setEnabled(False)
        self.btn_start.setText(self._t("ИДЕТ ОБРАБОТКА...", "PROCESSING..."))
        self.progress_bar_total.setValue(0)
        self.progress_bar_file.setValue(0)
        self._stage_start_time = 0.0
        self.detail_row.setVisible(True)
        self.btn_cancel.setEnabled(True)
        self.btn_cancel.setText(self._t("Отменить", "Cancel"))
        self.btn_cancel.setVisible(True)
        self.btn_open_result.setVisible(False)
        self._last_result_dir = self.output_dir or os.path.dirname(self.files_to_process[0])
        self.lbl_file_counter.setText(self._t(f"Файл 1 / {self.total_files}", f"File 1 / {self.total_files}"))
        self.lbl_current_file.setText("")
        self.lbl_stage.setText(self._t("●  Подготовка…", "●  Preparing…"))
        self.lbl_status.setText(self._t(f"Оценка: ~{estimate_str}", f"Estimate: ~{estimate_str}"))
        self._set_status(self._t(f"Обработка {self.total_files} файлов…", f"Processing {self.total_files} files…"))
        self.progress_timer.start(500)
        num_speakers = None
        if self.enable_diarization:
            value = self.entry_num_speakers.value()
            if value > 0:
                num_speakers = value
        snapshot = {
            "num_speakers": num_speakers,
            "enable_diarization": self.enable_diarization,
            "selected_formats": self._get_selected_formats(),
            "output_dir": self.output_dir,
            "files": list(self.files_to_process),
            "start_time": self.start_time,
        }
        self._set_processing_controls_enabled(False)
        threading.Thread(target=self._process_files, kwargs={"snapshot": snapshot}, daemon=True).start()

    def _set_processing_controls_enabled(self, enabled: bool):
        self.cb_diarization.setEnabled(enabled)
        self.entry_num_speakers.setEnabled(enabled and self.enable_diarization)
        self.btn_upload.setEnabled(enabled)
        self.input_path.setEnabled(enabled)
        self._update_files_controls()
        for cb in self.format_checkboxes.values():
            cb.setEnabled(enabled)
        if enabled:
            for fmt in ('txt_diarize', 'txt_diarize_timecodes'):
                cb = self.format_checkboxes.get(fmt)
                if cb:
                    cb.setEnabled(self.enable_diarization)

    def _process_files(self, snapshot: dict):
        num_speakers = snapshot["num_speakers"]
        enable_diarization = snapshot["enable_diarization"]
        selected_formats = snapshot["selected_formats"]
        output_dir = snapshot["output_dir"]
        files = snapshot["files"]
        start_time = snapshot["start_time"]
        total_files = len(files)
        try:
            if not self.model_loader.load_model(self.log):
                self.signals.processing_finished.emit(False, self._t("Не удалось загрузить модель", "Failed to load model"))
                return
            processor = TranscriptionProcessor(
                self.model_loader, self.stats, self.log,
                progress_callback=self._on_file_progress
            )
            if enable_diarization:
                self.log(self._t(f"Количество спикеров: {num_speakers if num_speakers else 'автоопределение'}", f"Speaker count: {num_speakers if num_speakers else 'auto-detect'}"))
            files_processed = 0
            files_failed = 0
            failed_names = []
            time_spent = 0.0
            generated_transcript_files = []
            for i, filepath in enumerate(files):
                if self._cancel_requested:
                    self.log(self._t("Обработка отменена пользователем", "Processing cancelled by user"))
                    break
                try:
                    self.current_file_start_time = time.time()
                    self.signals.current_file_info.emit(os.path.basename(filepath))
                    file_output_dir = output_dir if output_dir else os.path.dirname(filepath)
                    result = processor.process_file(
                        filepath, file_output_dir, i, total_files,
                        enable_diarization=enable_diarization,
                        num_speakers=num_speakers,
                        output_formats=selected_formats
                    )
                    self.stats.add_processing_record(
                        file_path=result['file_path'],
                        file_size=result['file_size'],
                        duration=result.get('media_duration', 0),
                        conversion_time=result['conversion_time'],
                        transcription_time=result['transcription_time'],
                        success=result['success']
                    )
                    if result['success']:
                        files_processed += 1
                        for saved_file in result.get('saved_files', []):
                            if saved_file.lower().endswith(('.txt', '.md', '.srt', '.vtt')):
                                generated_transcript_files.append(saved_file)
                    else:
                        files_failed += 1
                        failed_names.append(os.path.basename(filepath))
                    time_spent += result['total_time']
                except Exception as e:
                    files_failed += 1
                    failed_names.append(os.path.basename(filepath))
                    self.log(self._t(f"Ошибка при обработке файла {os.path.basename(filepath)}: {str(e)}", f"Error while processing file {os.path.basename(filepath)}: {str(e)}"))
                    continue
                finally:
                    self.files_processed = files_processed
                    self.time_spent = time_spent
            total_elapsed = time.time() - start_time
            self.log(self._t("=== ОБРАБОТКА ЗАВЕРШЕНА ===", "=== PROCESSING FINISHED ==="))
            self.log(self._t(f"Общее время обработки: {self.time_formatter.format_duration(total_elapsed)}", f"Total processing time: {self.time_formatter.format_duration(total_elapsed)}"))
            self.log(self._t(f"Успешно: {files_processed}/{total_files}" + (f", с ошибками: {files_failed}" if files_failed else ""), f"Successful: {files_processed}/{total_files}" + (f", with errors: {files_failed}" if files_failed else "")))
            cancelled = self._cancel_requested
            duration_str = self.time_formatter.format_duration(total_elapsed)
            if cancelled:
                message = self._t(f"Отменено. Обработано {files_processed}/{total_files} за {duration_str}", f"Cancelled. Processed {files_processed}/{total_files} in {duration_str}")
            elif files_failed:
                message = self._t(f"Готово с ошибками: {files_processed}/{total_files} успешно за {duration_str}", f"Completed with errors: {files_processed}/{total_files} successful in {duration_str}")
            else:
                message = self._t(f"Завершено за {duration_str}", f"Completed in {duration_str}")
            if failed_names:
                shown = ", ".join(failed_names[:5])
                if len(failed_names) > 5:
                    shown += self._t(f" и ещё {len(failed_names) - 5}", f" and {len(failed_names) - 5} more")
                message += self._t(f"\nНе удалось: {shown}\nПодробности — на вкладке «Журнал обработки».", f"\nFailed: {shown}\nSee details in the 'Processing log' tab.")
            success = (files_processed > 0) and (files_failed == 0) and not cancelled
            self._last_generated_transcript_files = generated_transcript_files
            self.signals.processing_finished.emit(success, message)
        except Exception as e:
            self.log(self._t(f"Критическая ошибка: {str(e)}", f"Critical error: {str(e)}"))
            self.signals.processing_finished.emit(False, self._t(f"Ошибка: {str(e)}", f"Error: {str(e)}"))

    # ──────────────────────────────────────────────────────────────
    # Прогресс
    # ──────────────────────────────────────────────────────────────

    _STAGE_NAMES = {'conversion': ('Конвертация…', 'Converting…'), 'transcription': ('Распознавание речи…', 'Speech recognition…')}

    def _on_file_progress(self, stage: str, progress: float):
        self.signals.stage_update.emit(stage, progress)

    def _on_stage_update(self, stage: str, progress: float):
        if stage != self.current_stage:
            self.current_stage = stage
            self._stage_start_time = time.time()
        self.current_stage_progress = progress
        self._refresh_progress()

    def _estimate_file_progress(self) -> float:
        if self.current_file_start_time <= 0 or not self.files_to_process:
            return 0.0
        idx = min(self.files_processed, len(self.files_to_process) - 1)
        total_est = self.file_estimates.get(self.files_to_process[idx], 30) or 30
        band = self._CONVERSION_BAND
        stage_elapsed = (time.time() - self._stage_start_time) if self._stage_start_time else 0.0
        if self.current_stage == 'conversion':
            conv_est = max(1.0, total_est * band)
            return band * min(1.0, stage_elapsed / conv_est)
        if self.current_stage == 'transcription':
            trans_est = max(1.0, total_est * (1 - band))
            return band + (0.99 - band) * min(1.0, stage_elapsed / trans_est)
        elapsed = time.time() - self.current_file_start_time
        return min(0.99, elapsed / total_est)

    def _refresh_progress(self):
        if not self.is_processing or self.total_files == 0 or not self.files_to_process:
            return
        files_done = min(self.files_processed, self.total_files)
        file_progress = self._estimate_file_progress()
        overall = min(0.99, (files_done + file_progress) / self.total_files)
        self.progress_bar_total.setValue(int(overall * 100))
        self.progress_bar_file.setValue(int(file_progress * 100))
        current_idx = min(files_done + 1, self.total_files)
        self.lbl_file_counter.setText(self._t(f"Файл {current_idx} / {self.total_files}", f"File {current_idx} / {self.total_files}"))
        stage_pair = self._STAGE_NAMES.get(self.current_stage or '', ('Подготовка…', 'Preparing…'))
        stage_name = stage_pair[0] if self._lang == 'ru' else stage_pair[1]
        self.lbl_stage.setText(f"●  {stage_name}  {int(file_progress * 100)}%")
        remaining_files = self.total_files - files_done
        if remaining_files > 0:
            if files_done > 0 and self.time_spent > 0:
                avg_time = self.time_spent / files_done
            else:
                avg_time = self.file_estimates.get(self.files_to_process[0], 30)
            self.lbl_status.setText(self._t(f"Осталось: ~{self.time_formatter.format_duration(avg_time * remaining_files)}", f"Remaining: ~{self.time_formatter.format_duration(avg_time * remaining_files)}"))

    def _update_progress_display(self):
        self._refresh_progress()

    def _update_total_progress(self, value: int):
        self.progress_bar_total.setValue(value)

    def _update_file_progress(self, value: int):
        self.progress_bar_file.setValue(value)

    def _update_current_file_info(self, info: str):
        self.current_stage = None
        self.current_stage_progress = 0.0
        display = info if len(info) <= 64 else f"…{info[-64:]}"
        self._current_filename = display
        self.lbl_current_file.setText(display)

    def _on_processing_finished(self, success: bool, message: str):
        self.is_processing = False
        self.progress_timer.stop()
        self.btn_start.setEnabled(True)
        self.btn_start.setText(self._t("ЗАПУСТИТЬ ОБРАБОТКУ", "START PROCESSING"))
        self.btn_cancel.setVisible(False)
        self.btn_cancel.setEnabled(True)
        self.btn_cancel.setText(self._t("Отменить", "Cancel"))
        self.progress_bar_total.setValue(100 if success else self.progress_bar_total.value())
        self.progress_bar_file.setValue(100 if success else self.progress_bar_file.value())
        self.lbl_stage.setText(self._t("✓  Готово", "✓  Done") if success else self._t("✕  Остановлено", "✕  Stopped"))
        self.lbl_status.setText(message.split("\n")[0])
        self._set_status(message.split("\n")[0])
        self._set_processing_controls_enabled(True)

        has_results = bool(self._last_result_dir) and os.path.isdir(self._last_result_dir)
        self.btn_open_result.setVisible(has_results)

        generated_files = [p for p in self._last_generated_transcript_files if os.path.isfile(p)]
        if generated_files:
            self.transcript_files_for_llm = generated_files
            self.user_settings.set_value("last_selected_transcript_files", generated_files)
            self.llm_transcript_dir = os.path.dirname(generated_files[0])
            if not self.llm_output_dir:
                self.llm_output_dir = self.llm_transcript_dir
                self._update_llm_output_dir_label(self.llm_output_dir)
            self.lbl_llm_files.setText(self._t(f"Выбрано транскриптов: {len(generated_files)}", f"Selected transcripts: {len(generated_files)}"))
            self.lbl_llm_files.setStyleSheet(self._transparent_label_style(self._colors()["text_sub"]))

        box = QMessageBox(self)
        box.setWindowTitle(self._t("Готово", "Done") if success else self._t("Завершено", "Finished"))
        box.setIcon(QMessageBox.Icon.Information if success else QMessageBox.Icon.Warning)
        box.setText((self._t("Обработка завершена!\n", "Processing completed!\n") + message) if success else message)
        open_btn = None
        if has_results:
            open_btn = box.addButton(self._t("Открыть папку с результатами", "Open results folder"), QMessageBox.ButtonRole.AcceptRole)
        box.addButton(QMessageBox.StandardButton.Ok)
        box.exec()
        if open_btn is not None and box.clickedButton() is open_btn:
            self._open_results_folder()

    def closeEvent(self, event):
        if self.is_processing or self.is_downloading:
            reply = QMessageBox.question(
                self,
                self._t("Внимание", "Attention"),
                self._t("Идёт обработка/загрузка. Закрыть приложение и прервать её?", "Processing/download is in progress. Close the app and interrupt it?"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No,
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return
            self._cancel_requested = True
            self.is_processing = False
        if self.output_dir:
            self.user_settings.set_last_output_dir(self.output_dir)
        if self.input_dir:
            self.user_settings.set_last_files_dir(self.input_dir)
        self._save_ui_settings()
        self._save_geometry()
        self.app_logger.log_session_end()
        event.accept()

    # ──────────────────────────────────────────────────────────────
    # Ошибки загрузки модели
    # ──────────────────────────────────────────────────────────────

    def _show_model_error(self, message: str):
        QMessageBox.warning(self, self._t("Ошибка загрузки", "Model loading error"), message)


def _argv_open_paths(argv: list) -> list:
    paths = []
    for arg in argv[1:]:
        if arg.startswith("-psn_"):
            continue
        path = os.path.abspath(os.path.expanduser(arg))
        if os.path.exists(path):
            paths.append(path)
    return paths


def _instance_lock_path() -> Path:
    return user_config_dir() / _INSTANCE_LOCK_NAME


def _open_requests_path() -> Path:
    return user_config_dir() / _OPEN_REQUESTS_NAME


def _try_acquire_instance_lock():
    lock_path = _instance_lock_path()
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    lock_file = lock_path.open("w", encoding="utf-8")
    if fcntl is None:
        return lock_file
    try:
        fcntl.flock(lock_file.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
    except BlockingIOError:
        lock_file.close()
        return None
    lock_file.write(str(os.getpid()))
    lock_file.flush()
    return lock_file


def _queue_open_request(paths: list) -> None:
    queue_path = _open_requests_path()
    queue_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {"paths": paths, "pid": os.getpid(), "time": time.time()}
    with queue_path.open("a", encoding="utf-8") as queue:
        queue.write(json.dumps(payload, ensure_ascii=False) + "\n")


def _install_open_request_poller(window: GigaTranscriberQtApp):
    queue_path = _open_requests_path()
    timer = QTimer(window)

    def poll_requests():
        if not queue_path.exists():
            return
        try:
            lines = queue_path.read_text(encoding="utf-8").splitlines()
            queue_path.write_text("", encoding="utf-8")
        except OSError as exc:
            window.log(f"Не удалось прочитать очередь open requests: {exc}")
            return
        for line in lines:
            if not line.strip():
                continue
            try:
                payload = json.loads(line)
            except json.JSONDecodeError:
                continue
            paths = payload.get("paths") or []
            if isinstance(paths, list):
                window.open_paths_from_system(paths, append=True)

    timer.timeout.connect(poll_requests)
    timer.start(500)
    window._open_request_poller = timer
    QTimer.singleShot(0, poll_requests)
    return timer


def run_qt_app(app=None):
    """Запускает приложение на PyQt6.

    app: уже созданный QApplication (используется на этапе выбора устройства).
    """
    if sys.platform == 'win32':
        try:
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('GigaAM.Transcriber.v3')
        except Exception:
            pass

    app = app or QApplication.instance() or GigaApplication(sys.argv)
    app.setFont(QFontDatabase.systemFont(QFontDatabase.SystemFont.GeneralFont))

    initial_open_paths = _argv_open_paths(sys.argv)
    instance_lock = getattr(app, "_gigaam_instance_lock_file", None) or _try_acquire_instance_lock()
    if instance_lock is None:
        _queue_open_request(initial_open_paths)
        sys.exit(0)

    icon_path = os.path.join(
        getattr(sys, '_MEIPASS', os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
        'icon.ico'
    )
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))

    window = GigaTranscriberQtApp()
    window._instance_lock_file = instance_lock
    _install_open_request_poller(window)
    if isinstance(app, GigaApplication):
        app.file_open_requested.connect(lambda paths: window.open_paths_from_system(paths, append=True))
    window.show()
    if initial_open_paths:
        QTimer.singleShot(0, lambda: window.open_paths_from_system(initial_open_paths, append=True))
    if isinstance(app, GigaApplication):
        pending_open_paths = app.take_pending_open_paths()
        if pending_open_paths:
            QTimer.singleShot(0, lambda paths=pending_open_paths: window.open_paths_from_system(paths, append=True))

    if sys.platform == 'win32' and os.path.exists(icon_path):
        def _set_win32_icon():
            try:
                import ctypes
                hwnd = int(window.winId())
                hicon = ctypes.windll.user32.LoadImageW(None, icon_path, 1, 0, 0, 0x10 | 0x40)
                if hicon:
                    ctypes.windll.user32.SendMessageW(hwnd, 0x0080, 0, hicon)
                    ctypes.windll.user32.SendMessageW(hwnd, 0x0080, 1, hicon)
            except Exception:
                pass
        QTimer.singleShot(100, _set_win32_icon)

    sys.exit(app.exec())


if __name__ == "__main__":
    run_qt_app()
